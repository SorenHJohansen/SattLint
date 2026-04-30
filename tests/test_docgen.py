from pathlib import Path
from typing import Any, cast

import pytest
from docx import Document
from docx.document import Document as DocClass
from openpyxl.worksheet.worksheet import Worksheet

from sattline_parser.grammar import constants as const
from sattline_parser.models.ast_model import (
    BasePicture,
    Equation,
    FrameModule,
    ModuleCode,
    ModuleHeader,
    ModuleTypeDef,
    ModuleTypeInstance,
    ParameterMapping,
    Sequence,
    SFCAlternative,
    SFCBreak,
    SFCCodeBlocks,
    SFCFork,
    SFCParallel,
    SFCStep,
    SFCSubsequence,
    SFCTransition,
    SFCTransitionSub,
    SingleModule,
    Variable,
)
from sattlint import config as config_module
from sattlint.docgenerator import configgen
from sattlint.docgenerator.classification import (
    DocumentationClassification,
    DocumentedModule,
    _collect_documented_modules,
    _descendants_of,
    _equals_pattern,
    _has_descendant_marker_match,
    _label_variants,
    _looks_like_equipment_module,
    _looks_like_operation,
    _looks_like_unit_root,
    _looks_like_wrapper_name,
    _marker_anchor,
    _matches_category_heuristic,
    _matches_rule,
    _normalize_requested_values,
    _resolve_instance_moduletype,
    _resolve_scope_paths,
    classify_documentation_structure,
    discover_documentation_unit_candidates,
    document_scope_summary,
)
from sattlint.docgenerator.docgen import (
    DocumentUnit,
    _calculation_rows,
    _communication_rows,
    _configurable_parameter_rows,
    _ensure_styles,
    _entry_variable,
    _entry_variable_text,
    _event_rows,
    _event_table_rows,
    _first_non_empty,
    _format_coord,
    _mapping_source_text,
    _mapping_target_name,
    _message_rows,
    _module_description,
    _pid_controller_rows,
    _prettify_name,
    _render_equipment_module_section,
    _sequence_render_rows,
    _sequence_rows,
    _simple_name_tag_rows,
    _special_logging_rows,
    _state_logic_summary,
    _state_rows,
    _value_text,
    _variable_rows,
    generate_docx,
)


def _hdr(name: str) -> ModuleHeader:
    return ModuleHeader(name=name, invoke_coord=(0.0, 0.0, 0.0, 1.0, 1.0))


def _literal_mapping(target: str, value: object) -> ParameterMapping:
    return ParameterMapping(
        target=target,
        source_type=const.KEY_VALUE,
        is_duration=False,
        is_source_global=False,
        source=None,
        source_literal=value,
    )


def _table_headers(document: DocClass) -> list[list[str]]:
    return [[cell.text.strip() for cell in table.rows[0].cells] for table in document.tables if table.rows]


def _table_text(document: DocClass) -> list[str]:
    return [
        cell.text.strip() for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip()
    ]


def _write_text(path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def _active_worksheet(workbook: Any) -> Worksheet:
    worksheet = workbook.active
    assert worksheet is not None
    return cast(Worksheet, worksheet)


def _typed_extractor(extractor: Any) -> configgen.SattLineConfigExtractor:
    return cast(configgen.SattLineConfigExtractor, extractor)


def _documented_instance(
    name: str,
    path: tuple[str, ...],
    *,
    moduletype_name: str,
    moduleparameters: list[Variable] | None = None,
    localvariables: list[Variable] | None = None,
    parametermappings: list[ParameterMapping] | None = None,
    modulecode: ModuleCode | None = None,
    resolved_moduletype: ModuleTypeDef | None = None,
) -> DocumentedModule:
    instance = ModuleTypeInstance(
        header=_hdr(name),
        moduletype_name=moduletype_name,
        parametermappings=list(parametermappings or []),
    )
    return DocumentedModule(
        node=instance,
        path=path,
        kind="moduletype_instance",
        current_library="ProjectLib",
        resolved_moduletype=resolved_moduletype,
        moduleparameters=tuple(moduleparameters or []),
        localvariables=tuple(localvariables or []),
        parametermappings=tuple(parametermappings or []),
        modulecode=modulecode,
    )


def _build_documentation_fixture() -> BasePicture:
    mes_state_control = ModuleTypeDef(name="MES_StateControl", origin_lib="NNEMESIFLib")
    equip_marker = ModuleTypeDef(name="EquipModCoordinate", origin_lib="nnestruct")
    recipe_param = ModuleTypeDef(name="RecParReal", origin_lib="NNELib")
    engineering_param = ModuleTypeDef(name="EngParReal", origin_lib="NNELib")
    user_param = ModuleTypeDef(name="UsrParReal", origin_lib="NNELib")

    operation_wrapper = ModuleTypeDef(
        name="OperationWrapper",
        origin_lib="ProjectLib",
        submodules=[
            ModuleTypeInstance(header=_hdr("MESInfo"), moduletype_name="MES_StateControl"),
            ModuleTypeInstance(header=_hdr("RecipeSP"), moduletype_name="RecParReal"),
            ModuleTypeInstance(header=_hdr("EngineeringLimit"), moduletype_name="EngParReal"),
            ModuleTypeInstance(header=_hdr("UserTarget"), moduletype_name="UsrParReal"),
        ],
        modulecode=ModuleCode(
            equations=[
                Equation(
                    name="Main",
                    position=(0.0, 0.0),
                    size=(1.0, 1.0),
                    code=["Ready = True"],
                )
            ]
        ),
    )

    equipment_wrapper = ModuleTypeDef(
        name="TankEM",
        origin_lib="ProjectLib",
        submodules=[
            ModuleTypeInstance(header=_hdr("Coordinate"), moduletype_name="EquipModCoordinate"),
            ModuleTypeInstance(header=_hdr("SpeedLimit"), moduletype_name="EngParReal"),
        ],
        modulecode=ModuleCode(
            equations=[
                Equation(
                    name="Physical",
                    position=(0.0, 0.0),
                    size=(1.0, 1.0),
                    code=["ValveOpen = True"],
                )
            ]
        ),
    )

    appl_tank = ModuleTypeDef(
        name="ApplTank",
        origin_lib="ProjectLib",
        submodules=[
            ModuleTypeInstance(header=_hdr("Prepare"), moduletype_name="OperationWrapper"),
            ModuleTypeInstance(header=_hdr("Tank"), moduletype_name="TankEM"),
        ],
    )

    dilute_unit = ModuleTypeDef(
        name="XDilute_221X251XY",
        origin_lib="ProjectLib",
        submodules=[
            ModuleTypeInstance(header=_hdr("DiluteTank"), moduletype_name="TankEM"),
        ],
    )

    return BasePicture(
        header=_hdr("BasePicture"),
        origin_lib="ProjectLib",
        moduletype_defs=[
            mes_state_control,
            equip_marker,
            recipe_param,
            engineering_param,
            user_param,
            operation_wrapper,
            equipment_wrapper,
            appl_tank,
            dilute_unit,
        ],
        submodules=[
            ModuleTypeInstance(header=_hdr("UnitA"), moduletype_name="ApplTank"),
            ModuleTypeInstance(header=_hdr("DilutionTrain"), moduletype_name="XDilute_221X251XY"),
            ModuleTypeInstance(header=_hdr("GlobalLimit"), moduletype_name="EngParReal"),
        ],
        library_dependencies={"projectlib": ["nnemesiflib", "nnestruct", "nnelib"]},
    )


def _build_wrapper_documentation_fixture() -> BasePicture:
    mes_state_control = ModuleTypeDef(name="MES_StateControl")
    equip_marker = ModuleTypeDef(name="EquipModCoordinate")
    state_logic = ModuleTypeDef(name="StateLogic")
    analog_input = ModuleTypeDef(name="AnalogInput")
    recipe_param = ModuleTypeDef(name="RecParReal")
    engineering_param = ModuleTypeDef(name="EngParReal")

    operation_l2 = ModuleTypeDef(
        name="OperationL2",
        submodules=[
            ModuleTypeInstance(header=_hdr("MES_StateControl"), moduletype_name="MES_StateControl"),
        ],
    )
    operation_l1 = ModuleTypeDef(
        name="OperationL1",
        submodules=[
            ModuleTypeInstance(header=_hdr("L2"), moduletype_name="OperationL2"),
        ],
    )
    operation_wrapper = ModuleTypeDef(
        name="OperationWrapper",
        submodules=[
            ModuleTypeInstance(header=_hdr("L1"), moduletype_name="OperationL1"),
            ModuleTypeInstance(header=_hdr("RecipeSP"), moduletype_name="RecParReal"),
            ModuleTypeInstance(header=_hdr("EngineeringLimit"), moduletype_name="EngParReal"),
        ],
    )
    opr_frame = ModuleTypeDef(
        name="OprFrameWrapper",
        submodules=[
            ModuleTypeInstance(
                header=_hdr("Prepare"),
                moduletype_name="OperationWrapper",
                parametermappings=[_literal_mapping("Name", "Prepare")],
            ),
        ],
    )
    operations = ModuleTypeDef(
        name="OperationsWrapper",
        submodules=[
            ModuleTypeInstance(header=_hdr("OprFrame"), moduletype_name="OprFrameWrapper"),
        ],
    )
    unit_control = ModuleTypeDef(
        name="UnitControlWrapper",
        submodules=[
            ModuleTypeInstance(header=_hdr("Operations"), moduletype_name="OperationsWrapper"),
        ],
    )

    equip_panel = ModuleTypeDef(
        name="EquipPanel",
        submodules=[
            ModuleTypeInstance(header=_hdr("Coordinate"), moduletype_name="EquipModCoordinate"),
            ModuleTypeInstance(header=_hdr("Stop"), moduletype_name="StateLogic"),
        ],
    )
    equip_display = ModuleTypeDef(
        name="EquipDisplay",
        submodules=[
            ModuleTypeInstance(header=_hdr("Panel"), moduletype_name="EquipPanel"),
        ],
    )
    equip_l2 = ModuleTypeDef(
        name="EquipL2",
        submodules=[
            ModuleTypeInstance(header=_hdr("Display"), moduletype_name="EquipDisplay"),
        ],
    )
    equip_l1 = ModuleTypeDef(
        name="EquipL1",
        submodules=[
            ModuleTypeInstance(header=_hdr("L2"), moduletype_name="EquipL2"),
        ],
    )
    tank_em = ModuleTypeDef(
        name="TankEM",
        moduleparameters=[Variable(name="Name", datatype="STRING")],
        submodules=[
            ModuleTypeInstance(header=_hdr("L1"), moduletype_name="EquipL1"),
            ModuleTypeInstance(header=_hdr("SpeedLimit"), moduletype_name="EngParReal"),
        ],
    )

    unit_type = ModuleTypeDef(
        name="UnitType",
        moduleparameters=[
            Variable(name="Name", datatype="STRING"),
            Variable(name="HeaderName", datatype="STRING"),
            Variable(name="SectionName", datatype="STRING"),
            Variable(name="InletPW", datatype="STRING"),
        ],
        submodules=[
            ModuleTypeInstance(
                header=_hdr("EM_FILL"),
                moduletype_name="TankEM",
                parametermappings=[_literal_mapping("Name", "Fill")],
            ),
            ModuleTypeInstance(header=_hdr("TT100"), moduletype_name="AnalogInput"),
            ModuleTypeInstance(header=_hdr("UnitControl"), moduletype_name="UnitControlWrapper"),
        ],
    )

    return BasePicture(
        header=_hdr("BasePicture"),
        moduletype_defs=[
            mes_state_control,
            equip_marker,
            state_logic,
            analog_input,
            recipe_param,
            engineering_param,
            operation_l2,
            operation_l1,
            operation_wrapper,
            opr_frame,
            operations,
            unit_control,
            equip_panel,
            equip_display,
            equip_l2,
            equip_l1,
            tank_em,
            unit_type,
        ],
        submodules=[
            ModuleTypeInstance(
                header=_hdr("UnitA"),
                moduletype_name="UnitType",
                parametermappings=[
                    _literal_mapping("Name", "238A"),
                    _literal_mapping("HeaderName", "UF/DF Tank"),
                    _literal_mapping("SectionName", "UF/DF"),
                    _literal_mapping("InletPW", "PWTransfer"),
                ],
            ),
        ],
    )


def _build_sequence_doc_fixture() -> BasePicture:
    mes_state_control = ModuleTypeDef(name="MES_StateControl")
    recipe_param = ModuleTypeDef(name="RecParReal")
    engineering_param = ModuleTypeDef(name="EngParReal")

    operation_wrapper = ModuleTypeDef(
        name="OperationWrapper",
        submodules=[
            ModuleTypeInstance(header=_hdr("MESInfo"), moduletype_name="MES_StateControl"),
            ModuleTypeInstance(header=_hdr("RecipeSP"), moduletype_name="RecParReal"),
            ModuleTypeInstance(header=_hdr("EngineeringLimit"), moduletype_name="EngParReal"),
        ],
        modulecode=ModuleCode(
            sequences=[
                Sequence(
                    name="MainSequence",
                    type="sequence",
                    position=(0.0, 0.0),
                    size=(1.0, 1.0),
                    code=[
                        SFCStep(
                            kind="init",
                            name="Init",
                            code=SFCCodeBlocks(exit=[("assign", {"var_name": "p.GotoRunDone"}, True)]),
                        ),
                        SFCTransition(
                            name="Tr1",
                            condition=(
                                "OR",
                                [
                                    {"var_name": "p.GotoRun"},
                                    {"var_name": "p.Run"},
                                ],
                            ),
                        ),
                        SFCStep(
                            kind="step",
                            name="Tarering",
                            code=SFCCodeBlocks(
                                enter=[
                                    (
                                        "FunctionCall",
                                        "CopyString",
                                        [
                                            {"var_name": "StepText.Tarering"},
                                            {"var_name": "Step"},
                                            {"var_name": "si"},
                                        ],
                                    )
                                ],
                                exit=[("assign", {"var_name": "Dv.Tara_execute"}, True)],
                            ),
                        ),
                        SFCTransition(
                            name="Tr2",
                            condition=(
                                "AND",
                                [
                                    ("NOT", {"var_name": "Dv.Tara_execute"}),
                                    ("NOT", {"var_name": "ProfibusError"}),
                                ],
                            ),
                        ),
                    ],
                )
            ]
        ),
    )

    unit_type = ModuleTypeDef(
        name="UnitType",
        moduleparameters=[
            Variable(name="Name", datatype="STRING"),
            Variable(name="HeaderName", datatype="STRING"),
            Variable(name="SectionName", datatype="STRING"),
        ],
        submodules=[
            ModuleTypeInstance(
                header=_hdr("Prepare"),
                moduletype_name="OperationWrapper",
                parametermappings=[_literal_mapping("Name", "Prepare")],
            ),
        ],
    )

    return BasePicture(
        header=_hdr("BasePicture"),
        moduletype_defs=[
            mes_state_control,
            recipe_param,
            engineering_param,
            operation_wrapper,
            unit_type,
        ],
        submodules=[
            ModuleTypeInstance(
                header=_hdr("UnitA"),
                moduletype_name="UnitType",
                parametermappings=[
                    _literal_mapping("Name", "238A"),
                    _literal_mapping("HeaderName", "UF/DF Tank"),
                    _literal_mapping("SectionName", "UF/DF"),
                ],
            ),
        ],
    )


def _build_empty_sequence_doc_fixture() -> BasePicture:
    mes_state_control = ModuleTypeDef(name="MES_StateControl")

    operation_wrapper = ModuleTypeDef(
        name="OperationWrapper",
        submodules=[
            ModuleTypeInstance(header=_hdr("MESInfo"), moduletype_name="MES_StateControl"),
        ],
        modulecode=ModuleCode(
            sequences=[
                Sequence(
                    name="EmptySequence",
                    type="sequence",
                    position=(0.0, 0.0),
                    size=(1.0, 1.0),
                    code=[],
                )
            ]
        ),
    )

    unit_type = ModuleTypeDef(
        name="UnitType",
        moduleparameters=[
            Variable(name="Name", datatype="STRING"),
            Variable(name="HeaderName", datatype="STRING"),
            Variable(name="SectionName", datatype="STRING"),
        ],
        submodules=[
            ModuleTypeInstance(
                header=_hdr("Prepare"),
                moduletype_name="OperationWrapper",
                parametermappings=[_literal_mapping("Name", "Prepare")],
            ),
        ],
    )

    return BasePicture(
        header=_hdr("BasePicture"),
        moduletype_defs=[mes_state_control, operation_wrapper, unit_type],
        submodules=[
            ModuleTypeInstance(
                header=_hdr("UnitA"),
                moduletype_name="UnitType",
                parametermappings=[
                    _literal_mapping("Name", "238A"),
                    _literal_mapping("HeaderName", "UF/DF Tank"),
                    _literal_mapping("SectionName", "UF/DF"),
                ],
            ),
        ],
    )


def _build_version_drift_doc_fixture() -> BasePicture:
    variant_a = SingleModule(
        header=_hdr("Mixer"),
        datecode=100,
        moduledef=None,
        moduleparameters=[],
        localvariables=[Variable(name="Output", datatype="INTEGER")],
        submodules=[],
        modulecode=ModuleCode(
            equations=[
                Equation(
                    name="Logic",
                    position=(0.0, 0.0),
                    size=(1.0, 1.0),
                    code=[("assign", {"var_name": "Output"}, 1)],
                )
            ],
            sequences=[],
        ),
        parametermappings=[],
    )
    variant_b = SingleModule(
        header=_hdr("Mixer"),
        datecode=200,
        moduledef=None,
        moduleparameters=[],
        localvariables=[Variable(name="Output", datatype="INTEGER")],
        submodules=[],
        modulecode=ModuleCode(
            equations=[
                Equation(
                    name="Logic",
                    position=(0.0, 0.0),
                    size=(1.0, 1.0),
                    code=[("assign", {"var_name": "Output"}, 2)],
                )
            ],
            sequences=[],
        ),
        parametermappings=[],
    )
    cast(Any, variant_a).origin_file = "BasePicture.s"
    cast(Any, variant_b).origin_file = "BasePicture.s"
    return BasePicture(
        header=_hdr("BasePicture"),
        origin_file="BasePicture.s",
        submodules=[variant_a, variant_b],
    )


def test_classify_documentation_structure_detects_recursive_categories():
    base_picture = _build_documentation_fixture()

    classification = classify_documentation_structure(
        base_picture,
        config_module.get_documentation_config(),
    )

    operations = {entry.name for entry in classification.categories["ops"]}
    equipment_modules = {entry.name for entry in classification.categories["em"]}
    recipe_parameters = {entry.name for entry in classification.categories["rp"]}
    engineering_parameters = {entry.name for entry in classification.categories["ep"]}
    user_parameters = {entry.name for entry in classification.categories["up"]}

    assert operations == {"Prepare"}
    assert equipment_modules == {"Tank", "DiluteTank"}
    assert "RecipeSP" in recipe_parameters
    assert "EngineeringLimit" in engineering_parameters
    assert "GlobalLimit" in engineering_parameters
    assert "UserTarget" in user_parameters

    operation_entry = classification.categories["ops"][0]
    assert {entry.name for entry in classification.descendants(operation_entry, category="rp")} == {"RecipeSP"}


def test_generate_docx_renders_fs_style_sections(tmp_path):
    base_picture = _build_documentation_fixture()
    out_path = tmp_path / "functional-spec.docx"

    generate_docx(
        base_picture,
        out_path,
        documentation_config=config_module.get_documentation_config(),
    )

    assert out_path.exists()

    document = Document(out_path)
    paragraph_text = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    assert "Introduction" in paragraph_text
    assert "S88 Physical model" in paragraph_text
    assert "S88 Procedural model" in paragraph_text
    assert "Measurements and logging" in paragraph_text
    assert "Communication" in paragraph_text
    assert "Equipment Module Tank" in paragraph_text
    assert "Operation Prepare" in paragraph_text
    assert "Change Log" in paragraph_text


def test_classify_documentation_scope_by_moduletype_name():
    base_picture = _build_documentation_fixture()
    documentation_cfg = config_module.get_documentation_config()
    documentation_cfg["units"] = {
        "mode": "moduletype_names",
        "instance_paths": [],
        "moduletype_names": ["ApplTank"],
    }

    classification = classify_documentation_structure(base_picture, documentation_cfg)

    assert classification.scope is not None
    assert [entry.name for entry in classification.scope.roots or []] == ["UnitA"]
    assert {entry.name for entry in classification.categories["ops"]} == {"Prepare"}
    assert {entry.name for entry in classification.categories["em"]} == {"Tank"}


def test_discover_documentation_unit_candidates_returns_unit_roots():
    base_picture = _build_documentation_fixture()
    classification = classify_documentation_structure(
        base_picture,
        config_module.get_documentation_config(),
    )

    candidates = discover_documentation_unit_candidates(classification)

    assert {entry.name for entry in candidates} == {"UnitA", "DilutionTrain"}


def test_wrapper_heavy_classification_anchors_real_sections():
    base_picture = _build_wrapper_documentation_fixture()

    classification = classify_documentation_structure(
        base_picture,
        config_module.get_documentation_config(),
    )

    assert {entry.name for entry in classification.categories["em"]} == {"EM_FILL"}
    assert {entry.name for entry in classification.categories["ops"]} == {"Prepare"}

    candidates = discover_documentation_unit_candidates(classification)
    assert {entry.name for entry in candidates} == {"UnitA"}


def test_generate_docx_renders_wrapper_heavy_fs_sections(tmp_path):
    base_picture = _build_wrapper_documentation_fixture()
    out_path = tmp_path / "functional-spec-wrapper.docx"

    generate_docx(
        base_picture,
        out_path,
        documentation_config=config_module.get_documentation_config(),
    )

    document = Document(out_path)
    paragraph_text = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    assert "Process Cell - UF/DF" in paragraph_text
    assert "Unit Class definition: UF/DF Tank" in paragraph_text
    assert "Measurements and logging" in paragraph_text
    assert "Communication" in paragraph_text
    assert "Equipment Module Fill" in paragraph_text
    assert "Operation Prepare" in paragraph_text
    assert "Change Log" in paragraph_text


def test_generate_docx_respects_single_unit_scope(tmp_path):
    base_picture = _build_documentation_fixture()
    out_path = tmp_path / "functional-spec-scoped.docx"
    documentation_cfg = config_module.get_documentation_config()
    documentation_cfg["units"] = {
        "mode": "moduletype_names",
        "instance_paths": [],
        "moduletype_names": ["ApplTank"],
    }

    generate_docx(
        base_picture,
        out_path,
        documentation_config=documentation_cfg,
    )

    document = Document(out_path)
    paragraph_text = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    assert "Operation Prepare" in paragraph_text
    assert "Equipment Module Tank" in paragraph_text
    assert "DilutionTrain" not in paragraph_text
    assert "Equipment Module DiluteTank" not in paragraph_text


def test_generate_docx_uses_template_shaped_front_matter_and_tables(tmp_path):
    base_picture = _build_wrapper_documentation_fixture()
    out_path = tmp_path / "functional-spec-template-shaped.docx"

    generate_docx(
        base_picture,
        out_path,
        documentation_config=config_module.get_documentation_config(),
    )

    document = Document(out_path)
    headers = _table_headers(document)

    assert ["NNE Author", "NNE Author", "NNE Author"] in headers
    assert ["", "Document", "NN Doc. no."] in headers
    assert ["Unit", "Unit Class", "Danish Description", "Unit Definition"] in headers
    assert [
        "Measurement",
        "Tag",
        "Min",
        "Max",
        "Eng. Unit",
        "Log interval\n(Max)",
        "Dead-band\nrelative",
        "Log interval \n(Min)",
    ] in headers
    assert ["From", "Comment"] in headers


def test_generate_docx_renders_upgrade_insights_for_version_drift(tmp_path):
    base_picture = _build_version_drift_doc_fixture()
    out_path = tmp_path / "functional-spec-upgrades.docx"

    generate_docx(
        base_picture,
        out_path,
        documentation_config=config_module.get_documentation_config(),
    )

    document = Document(out_path)
    paragraph_text = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    assert "Upgrade insights" in paragraph_text
    assert "Module Mixer" in paragraph_text
    assert any("Equation 'Logic' changed" in paragraph for paragraph in paragraph_text)


def test_generate_docx_renders_sfc_sequences_as_tables(tmp_path):
    base_picture = _build_sequence_doc_fixture()
    out_path = tmp_path / "functional-spec-sfc.docx"

    generate_docx(
        base_picture,
        out_path,
        documentation_config=config_module.get_documentation_config(),
    )

    document = Document(out_path)
    headers = _table_headers(document)
    table_text = _table_text(document)

    assert ["Type", "Name", "Condition / Detail", "Enter", "Active", "Exit"] in headers
    assert "Sub sequence - MainSequence" in [
        paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
    ]
    assert "Init step" in table_text
    assert "Init" in table_text
    assert "p.GotoRunDone = True" in table_text
    assert "Transition" in table_text
    assert "Tr1" in table_text
    assert "p.GotoRun OR \np.Run" in table_text
    assert "Step" in table_text
    assert "Tarering" in table_text
    assert "CopyString(StepText.Tarering, Step, si)" in table_text
    assert "NOT(Dv.Tara_execute) AND \nNOT(ProfibusError)" in table_text
    assert not any("SFCStep(" in text for text in table_text)
    assert not any("SFCTransition(" in text for text in table_text)


def test_classify_documentation_scope_by_instance_path_deduplicates_and_tracks_unmatched_values():
    base_picture = _build_documentation_fixture()
    documentation_cfg = config_module.get_documentation_config()
    documentation_cfg["units"] = {
        "mode": "instance_paths",
        "instance_paths": ["BasePicture.UnitA", "UnitA", "Missing.Unit"],
        "moduletype_names": [],
    }

    classification = classify_documentation_structure(base_picture, documentation_cfg)

    assert classification.scope is not None
    assert classification.scope.mode == "instance_paths"
    assert classification.scope.requested_values == ["BasePicture.UnitA", "UnitA", "Missing.Unit"]
    assert [entry.name for entry in classification.scope.roots or []] == ["UnitA"]
    assert classification.scope.unmatched_values == ["Missing.Unit"]
    assert {entry.name for entry in classification.categories["ops"]} == {"Prepare"}
    assert {entry.name for entry in classification.categories["em"]} == {"Tank"}
    assert {entry.name for entry in classification.uncategorized} == {"MESInfo", "Coordinate"}


def test_classify_documentation_scope_falls_back_to_all_for_invalid_units_settings():
    base_picture = _build_documentation_fixture()

    invalid_shape = classify_documentation_structure(base_picture, {"units": "bad"})
    unknown_mode = classify_documentation_structure(base_picture, {"units": {"mode": "mystery"}})

    assert invalid_shape.scope is not None
    assert invalid_shape.scope.mode == "all"
    assert invalid_shape.scope.roots == []
    assert invalid_shape.scope.unmatched_values is None
    assert unknown_mode.scope is not None
    assert unknown_mode.scope.mode == "all"
    assert unknown_mode.scope.roots == []


def test_sequence_render_rows_cover_branching_nested_and_fallback_nodes():
    sequence = Sequence(
        name="ComplexSequence",
        type="sequence",
        position=(0.0, 0.0),
        size=(1.0, 1.0),
        code=[
            SFCAlternative(
                branches=[
                    [SFCStep(kind="init", name="AltInit", code=SFCCodeBlocks())],
                    [SFCBreak()],
                ]
            ),
            SFCParallel(
                branches=[
                    [SFCTransition(name="ParallelGate", condition=True)],
                    [SFCFork(target="ParallelDone")],
                ]
            ),
            SFCSubsequence(
                name="NestedSequence",
                body=[SFCStep(kind="step", name="NestedStep", code=SFCCodeBlocks())],
            ),
            SFCTransitionSub(
                name="TransitionBranch",
                body=[SFCTransition(name="NestedTransition", condition=False)],
            ),
            SFCFork(target="Elsewhere"),
            SFCBreak(),
            (const.KEY_ASSIGN, {"var_name": "Output"}, True),
        ],
    )

    rows = _sequence_render_rows(sequence)

    assert [row.node_type for row in rows] == [
        "Alternative",
        "Branch",
        "Init step",
        "Branch",
        "Break",
        "Parallel",
        "Branch",
        "Transition",
        "Branch",
        "Fork",
        "Subsequence",
        "Step",
        "Transition section",
        "Transition",
        "Fork",
        "Break",
        "Statement",
    ]
    assert rows[2].detail == "Alternative branch 1: Initial step"
    assert rows[7].detail == "Parallel branch 1: True"
    assert rows[11].detail == "Subsequence NestedSequence: Execution step"
    assert rows[13].detail == "Transition section TransitionBranch: False"
    assert rows[14].detail == "Fork target"
    assert rows[15].detail == "Break sequence flow"
    assert rows[16].detail == "Output = True"


def test_generate_docx_renders_appendix_for_uncategorized_modules(tmp_path):
    base_picture = _build_documentation_fixture()
    out_path = tmp_path / "functional-spec-appendix.docx"

    generate_docx(
        base_picture,
        out_path,
        documentation_config=config_module.get_documentation_config(),
    )

    document = Document(out_path)
    paragraph_text = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text = _table_text(document)

    assert "Appendix: Supporting modules" in paragraph_text
    assert "Other module instances" in paragraph_text
    assert "UnitA" in table_text
    assert "DilutionTrain" in table_text
    assert "GlobalLimit" not in table_text


def test_generate_docx_renders_empty_sequences_as_no_explicit_statements(tmp_path):
    base_picture = _build_empty_sequence_doc_fixture()
    out_path = tmp_path / "functional-spec-empty-sequence.docx"

    generate_docx(
        base_picture,
        out_path,
        documentation_config=config_module.get_documentation_config(),
    )

    document = Document(out_path)
    paragraph_text = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    assert "Sub sequence - EmptySequence" in paragraph_text
    assert "No explicit sequence statements detected." in paragraph_text


def test_classification_private_helpers_cover_scope_lookup_and_anchor_edges():
    unresolved_root = _documented_instance("Rootless", (), moduletype_name="LooseType")
    scoped_root = _documented_instance(
        "UnitAlpha",
        ("BasePicture", "AreaA", "UnitAlpha"),
        moduletype_name="UnitType",
        moduleparameters=[Variable(name="SectionName", datatype="STRING")],
    )
    wrapper = _documented_instance(
        "Operations",
        ("BasePicture", "AreaA", "UnitAlpha", "Operations"),
        moduletype_name="OperationsWrapper",
    )
    marker = _documented_instance(
        "Prepare",
        ("BasePicture", "AreaA", "UnitAlpha", "Operations", "OprFrame", "Prepare"),
        moduletype_name="OperationPhase",
    )
    entries = [scoped_root, wrapper, marker]
    entry_lookup = {entry.path: entry for entry in entries}

    assert unresolved_root.moduletype_name == "LooseType"
    assert unresolved_root.moduletype_label == "LooseType"
    assert unresolved_root.short_path == ""
    assert scoped_root.short_path == "AreaA.UnitAlpha"
    assert _matches_rule(scoped_root, {}, entries) is False
    assert _has_descendant_marker_match(scoped_root, entries, name_contains=["operation"], label_equals=[])
    assert _marker_anchor(marker, entry_lookup) is scoped_root

    roots, unmatched = _resolve_scope_paths([scoped_root], ["UnitAlpha", "Missing.Unit"])

    assert roots == [scoped_root]
    assert unmatched == ["Missing.Unit"]


def test_discover_documentation_unit_candidates_cover_fallback_paths():
    fallback_root = _documented_instance(
        "DeepUnit",
        ("BasePicture", "AreaA", "CellA", "DeepUnit"),
        moduletype_name="UnitType",
    )
    fallback_equipment = _documented_instance(
        "Pump",
        ("BasePicture", "AreaA", "CellA", "DeepUnit", "Pump"),
        moduletype_name="EquipModPump",
    )
    fallback_classification = DocumentationClassification(
        section_order=["em", "ops", "rp", "ep", "up"],
        categories={"em": [fallback_equipment], "ops": [], "rp": [], "ep": [], "up": []},
        uncategorized=[fallback_root],
        all_entries=[fallback_root, fallback_equipment],
    )

    top_level_root = _documented_instance(
        "LooseUnit",
        ("BasePicture", "LooseUnit"),
        moduletype_name="UnitType",
    )
    terminal_classification = DocumentationClassification(
        section_order=["em", "ops", "rp", "ep", "up"],
        categories={"em": [], "ops": [], "rp": [], "ep": [], "up": []},
        uncategorized=[top_level_root],
        all_entries=[top_level_root],
    )

    assert discover_documentation_unit_candidates(fallback_classification) == [fallback_root]
    assert discover_documentation_unit_candidates(terminal_classification) == [top_level_root]


def test_docgen_private_helpers_cover_mapping_and_support_row_filters():
    inlet_mapping = ParameterMapping(
        target={"var_name": "InletPW"},
        source_type=const.KEY_VALUE,
        is_duration=False,
        is_source_global=False,
        source={"var_name": "Feed.Line"},
        source_literal=None,
    )
    outlet_mapping = _literal_mapping("OutletPW", "Drain.Header")
    blank_mapping = _literal_mapping("Spare", None)

    equipment = _documented_instance(
        "Tank",
        ("BasePicture", "UnitA", "Tank"),
        moduletype_name="TankEM",
        localvariables=[Variable(name="Description", datatype="STRING", description="Tank description")],
        parametermappings=[inlet_mapping, outlet_mapping, blank_mapping],
    )
    state_panel = _documented_instance(
        "StateLamp",
        ("BasePicture", "UnitA", "Tank", "Panel", "StateLamp"),
        moduletype_name="StateDisplay",
    )
    toggle_panel = _documented_instance(
        "KaHCToggleStart",
        ("BasePicture", "UnitA", "Tank", "Panel", "KaHCToggleStart"),
        moduletype_name="ToggleWidget",
    )
    state_logic = _documented_instance(
        "StopLogic",
        ("BasePicture", "UnitA", "Tank", "Panel", "StateLamp", "StopLogic"),
        moduletype_name="StateLogic",
        modulecode=ModuleCode(
            equations=[Equation(name="Main", position=(0.0, 0.0), size=(1.0, 1.0), code=[True])],
            sequences=[
                Sequence(
                    name="StateSeq",
                    type="sequence",
                    position=(0.0, 0.0),
                    size=(1.0, 1.0),
                    code=[],
                )
            ],
        ),
    )
    message = _documented_instance(
        "OperatorMessage",
        ("BasePicture", "UnitA", "Tank", "OperatorMessage"),
        moduletype_name="OpMessageModule",
    )
    event = _documented_instance(
        "BatchJournal",
        ("BasePicture", "UnitA", "Tank", "BatchJournal"),
        moduletype_name="JournalModule",
    )
    special_log = _documented_instance(
        "AuditLog",
        ("BasePicture", "UnitA", "Tank", "AuditLog"),
        moduletype_name="LoggerModule",
    )
    classification = DocumentationClassification(
        section_order=["em", "ops", "rp", "ep", "up"],
        categories={"em": [], "ops": [], "rp": [], "ep": [], "up": []},
        uncategorized=[],
        all_entries=[equipment, state_panel, toggle_panel, state_logic, message, event, special_log],
    )
    unit = DocumentUnit(
        root=equipment,
        unit_code="UnitA",
        title="Unit A",
        unit_class="UnitType",
        section_name="Unit A",
        equipment_modules=[equipment],
        operations=[],
        recipe_parameters=[],
        engineering_parameters=[],
        user_parameters=[],
    )

    assert _mapping_target_name(inlet_mapping) == "InletPW"
    assert _mapping_source_text(inlet_mapping) == "Feed.Line"
    assert _mapping_source_text(outlet_mapping) == "Drain.Header"
    assert _mapping_source_text(blank_mapping) == ""
    assert _module_description(equipment) == "Tank description"
    assert _communication_rows(unit, direction="from") == [["Feed.Line", "Mapped via InletPW"]]
    assert _communication_rows(unit, direction="to") == [["Drain.Header", "Mapped via OutletPW"]]

    state_rows = _state_rows(equipment, classification)
    assert state_rows == [
        ["StateLamp", "Detected StateDisplay instance.", "1 state-logic modules, 1 sequences, 1 equation blocks"]
    ]
    assert _state_logic_summary(message, classification) == "Display-only state definition"
    assert [row[0] for row in _message_rows(equipment, classification)] == ["OperatorMessage"]
    assert [row[0] for row in _event_rows(equipment, classification)] == ["BatchJournal"]
    assert [row[0] for row in _special_logging_rows(equipment, classification)] == [
        "StopLogic",
        "BatchJournal",
        "AuditLog",
    ]


def test_classification_private_helpers_cover_collection_resolution_and_wrapper_variants():
    loop_type = ModuleTypeDef(
        name="LoopType",
        origin_lib="ProjectLib",
        submodules=[ModuleTypeInstance(header=_hdr("Self"), moduletype_name="LoopType")],
    )
    frame = FrameModule(header=_hdr("FrameA"), submodules=[], modulecode=None)
    unknown_instance = ModuleTypeInstance(header=_hdr("Unknown"), moduletype_name="MissingType")
    loop_instance = ModuleTypeInstance(header=_hdr("Loop"), moduletype_name="LoopType")
    base_picture = BasePicture(
        header=_hdr("BasePicture"),
        origin_lib="ProjectLib",
        moduletype_defs=[loop_type],
        submodules=[frame, unknown_instance, loop_instance],
    )

    entries = _collect_documented_modules(base_picture)
    names = [entry.name for entry in entries]

    assert names == ["FrameA", "Unknown", "Loop", "Self"]
    assert (
        _resolve_instance_moduletype(
            base_picture, unknown_instance, current_library="ProjectLib", unavailable_libraries=None
        )
        is None
    )
    assert [entry.name for entry in _descendants_of(entries[2], entries)] == ["Self"]
    assert _normalize_requested_values("bad") == []
    assert _looks_like_wrapper_name("Prefix:Panel") is True
    assert _label_variants("Parent:Child") == {"parent:child", "child"}
    assert _looks_like_equipment_module(
        _documented_instance("EM_Pump", ("BasePicture", "EM_Pump"), moduletype_name="PumpType")
    )
    assert _looks_like_operation(
        _documented_instance("Prepare", ("BasePicture", "OprFrame", "Prepare"), moduletype_name="OperationPhase")
    )
    assert not _looks_like_operation(
        _documented_instance("MES_Stop", ("BasePicture", "OprFrame", "MES_Stop"), moduletype_name="OperationPhase")
    )


def test_docgen_private_helpers_cover_scalar_and_renderer_fallback_paths(tmp_path):
    plain_source_mapping = ParameterMapping(
        target="Target",
        source_type=const.KEY_VALUE,
        is_duration=False,
        is_source_global=False,
        source=cast(Any, "Upstream.Tag"),
        source_literal=None,
    )
    ignored_target = _literal_mapping("ProgramName", "IgnoreMe")
    empty_target = _literal_mapping("", "IgnoreMe")
    configurable_target = _literal_mapping("TargetTag", "Mapped.Value")
    equipment = _documented_instance(
        "Valve",
        ("BasePicture", "UnitA", "Valve"),
        moduletype_name="ValveType",
        moduleparameters=[Variable(name="Name", datatype="STRING"), Variable(name="Speed", datatype="REAL")],
        localvariables=[Variable(name="LocalText", datatype="STRING", description="Fallback description")],
        parametermappings=[plain_source_mapping, ignored_target, empty_target, configurable_target],
    )
    unit = DocumentUnit(
        root=equipment,
        unit_code="UnitA",
        title="Unit A",
        unit_class="UnitType",
        section_name="Unit A",
        equipment_modules=[equipment],
        operations=[],
        recipe_parameters=[],
        engineering_parameters=[],
        user_parameters=[],
    )
    document = Document()
    _ensure_styles(document)
    classification = DocumentationClassification(
        section_order=["em", "ops", "rp", "ep", "up"],
        categories={"em": [], "ops": [], "rp": [], "ep": [], "up": []},
        uncategorized=[],
        all_entries=[equipment],
    )

    assert _value_text(None) == ""
    assert _value_text(5) == "5"
    assert _format_coord(None) == "<none>"
    assert _format_coord((1.0, 2.5)) == "(1, 2.5)"
    assert _prettify_name("") == ""
    assert _prettify_name("ALLCAPS") == "ALLCAPS"
    assert _mapping_source_text(plain_source_mapping) == "Upstream.Tag"
    assert _first_non_empty("", "Value") == "Value"
    assert _entry_variable(equipment, "speed") is not None
    assert _entry_variable_text(equipment, "LocalText") == "Fallback description"
    assert _configurable_parameter_rows(unit) == [
        ["Target", "", "UnitA", "Upstream.Tag"],
        ["TargetTag", "", "UnitA", "Mapped.Value"],
    ]
    assert _variable_rows([Variable(name="Speed", datatype="REAL", description="Valve speed")]) == [
        ["Speed", "real", "", "Valve speed"]
    ]

    _render_equipment_module_section(document, equipment, classification)

    assert _module_description(equipment) == "Fallback description"
    assert ["Tag", "Datatype", "Init", "Description"] in _table_headers(document)

    base_picture = _build_documentation_fixture()
    cast(Any, base_picture.submodules[0]).parametermappings = [
        _literal_mapping("OutletPW", "Downstream.Header"),
    ]
    out_path = tmp_path / "functional-spec-outbound.docx"

    generate_docx(
        base_picture,
        out_path,
        documentation_config=config_module.get_documentation_config(),
    )

    rendered = Document(out_path)
    assert ["To", "Comment"] in _table_headers(rendered)


def test_classification_and_docgen_final_helper_edges():
    top_level = _documented_instance("BasePicture", ("BasePicture",), moduletype_name="TopType")
    scoped_root = _documented_instance(
        "ScopedUnit",
        ("BasePicture", "ScopedUnit"),
        moduletype_name="UnitType",
        moduleparameters=[Variable(name="SectionName", datatype="STRING")],
    )
    shared_sequence = Sequence(name="Shared", type="sequence", position=(0.0, 0.0), size=(1.0, 1.0), code=[])
    equipment = _documented_instance(
        "Pump",
        ("BasePicture", "ScopedUnit", "Pump"),
        moduletype_name="EquipModPump",
        modulecode=ModuleCode(
            sequences=[shared_sequence],
        ),
    )
    duplicate_equipment = _documented_instance(
        "Pump",
        ("BasePicture", "ScopedUnit", "Pump"),
        moduletype_name="EquipModPump",
        modulecode=ModuleCode(
            sequences=[shared_sequence],
        ),
    )
    wrapper_parent = _documented_instance(
        "OprFrame",
        ("BasePicture", "Operations", "OprFrame"),
        moduletype_name="OprFrame",
    )
    wrapper_marker = _documented_instance(
        "Prepare",
        ("BasePicture", "Operations", "OprFrame", "Prepare"),
        moduletype_name="OperationPhase",
    )
    rootless = _documented_instance("Root", (), moduletype_name="RootType")
    short_descendant = _documented_instance("Solo", ("Solo",), moduletype_name="Display")
    off_panel_descendant = _documented_instance(
        "Lamp",
        ("BasePicture", "ScopedUnit", "Pump", "Display", "Lamp"),
        moduletype_name="Display",
    )
    toggle_descendant = _documented_instance(
        "KaHCToggleLamp",
        ("BasePicture", "ScopedUnit", "Pump", "Panel", "KaHCToggleLamp"),
        moduletype_name="Display",
    )
    state_descendant = _documented_instance(
        "StateLamp",
        ("BasePicture", "ScopedUnit", "Pump", "Panel", "StateLamp"),
        moduletype_name="Display",
    )
    duplicate_state_descendant = _documented_instance(
        "StateLamp",
        ("BasePicture", "ScopedUnit", "Pump", "Panel", "StateLamp"),
        moduletype_name="Display",
    )
    event_entry = _documented_instance(
        "Event01",
        ("BasePicture", "ScopedUnit", "Pump", "Event01"),
        moduletype_name="EventModule",
        parametermappings=[_literal_mapping("Severity", "High"), _literal_mapping("Activation", "WhenTrue")],
    )
    calc_entry = _documented_instance(
        "Calc01",
        ("BasePicture", "ScopedUnit", "Pump", "Calc01"),
        moduletype_name="CalcModule",
    )
    pid_entry = _documented_instance(
        "PC101",
        ("BasePicture", "ScopedUnit", "Pump", "PC101"),
        moduletype_name="ControlModule",
    )
    blank_source_unit = DocumentUnit(
        root=_documented_instance(
            "RootUnit",
            ("BasePicture", "RootUnit"),
            moduletype_name="UnitType",
            parametermappings=[_literal_mapping("OutletPW", None)],
        ),
        unit_code="RootUnit",
        title="Root Unit",
        unit_class="UnitType",
        section_name="Root Unit",
        equipment_modules=[],
        operations=[],
        recipe_parameters=[],
        engineering_parameters=[],
        user_parameters=[],
    )
    blank_target_unit = DocumentUnit(
        root=_documented_instance(
            "BlankTargetUnit",
            ("BasePicture", "BlankTargetUnit"),
            moduletype_name="UnitType",
            parametermappings=[_literal_mapping("", "PipeA")],
        ),
        unit_code="BlankTargetUnit",
        title="Blank Target Unit",
        unit_class="UnitType",
        section_name="Blank Target Unit",
        equipment_modules=[],
        operations=[],
        recipe_parameters=[],
        engineering_parameters=[],
        user_parameters=[],
    )
    classification = DocumentationClassification(
        section_order=["em", "ops", "rp", "ep", "up"],
        categories={"em": [equipment], "ops": [], "rp": [], "ep": [], "up": []},
        uncategorized=[scoped_root],
        all_entries=[
            scoped_root,
            equipment,
            duplicate_equipment,
            short_descendant,
            off_panel_descendant,
            toggle_descendant,
            state_descendant,
            duplicate_state_descendant,
            event_entry,
            calc_entry,
            pid_entry,
        ],
    )

    assert top_level.short_path == "BasePicture"
    assert document_scope_summary(scoped_root, classification) == {"ops": 0, "em": 1, "rp": 0, "ep": 0, "up": 0}
    assert _equals_pattern("UnitType", ["", "UnitType"]) is True
    assert _matches_rule(scoped_root, {"name_contains": ["Unit"]}, [scoped_root]) is True
    assert (
        _marker_anchor(wrapper_marker, {wrapper_parent.path: wrapper_parent, wrapper_marker.path: wrapper_marker})
        is wrapper_parent
    )
    assert _matches_category_heuristic(scoped_root, "other") is False
    assert not _looks_like_equipment_module(
        _documented_instance("Coordinate", ("BasePicture", "Coordinate"), moduletype_name="PumpType")
    )
    assert not _looks_like_operation(_documented_instance("Step", ("Step",), moduletype_name="OperationPhase"))
    assert _looks_like_unit_root(scoped_root, classification, categorized_paths=set()) is True
    assert _looks_like_wrapper_name("Prefix:Panel") is True

    assert _configurable_parameter_rows(blank_source_unit) == []
    assert _simple_name_tag_rows([calc_entry], "description") == [["Detected CalcModule instance.", "Calc01"]]
    assert _event_table_rows([event_entry]) == [["Event01", "Detected EventModule instance.", "High", "WhenTrue"]]
    assert _calculation_rows([calc_entry]) == [["Calc01", "Detected CalcModule instance.", "ScopedUnit.Pump.Calc01"]]
    assert _communication_rows(blank_target_unit, direction="to") == []
    assert _state_rows(rootless, classification) == [
        ["StateLamp", "Detected Display instance.", "Display-only state definition"]
    ]
    assert _pid_controller_rows(scoped_root, classification) == [
        ["PC101", "ControlModule", "ScopedUnit.Pump.PC101", "Detected ControlModule instance."]
    ]
    assert _sequence_rows(scoped_root, classification) == [("Shared", _sequence_render_rows(shared_sequence))]


def test_configgen_parse_configuration_file_extracts_programs_and_libraries(tmp_path):
    config_file = tmp_path / "Configuration" / "KaGC_Allf.k"
    _write_text(
        config_file,
        'Configuration ( Version "v1" Date "2026-04-29" Name "IgnoredByParser" )\n'
        'Program ( Name "MainProgram.z" Directory "UnitLib" MainProgram True )\n'
        'Program ( Name "SupportProgram" Directory "UnitLib" MainProgram False )\n'
        'Library ( Name "CommonLib.z" Directory "ProjectLib" )\n',
    )

    parser = configgen.ConfigurationFileParser()

    info = parser.parse_configuration_file(config_file)

    assert info is not None
    assert info.config_name == "KaGC_Allf"
    assert info.version == "v1"
    assert info.date == "2026-04-29"
    assert info.main_program == "MainProgram"
    assert info.programs == [
        {"name": "MainProgram", "directory": "UnitLib", "main_program": True},
        {"name": "SupportProgram", "directory": "UnitLib", "main_program": False},
    ]
    assert info.libraries == [{"name": "CommonLib", "directory": "ProjectLib"}]


def test_configgen_parse_configuration_file_returns_none_without_header(tmp_path):
    config_file = tmp_path / "Configuration" / "Broken.k"
    _write_text(config_file, 'Program ( Name "MainProgram.z" Directory "UnitLib" MainProgram True )\n')

    parser = configgen.ConfigurationFileParser()

    assert parser.parse_configuration_file(config_file) is None


def test_configgen_extractor_get_component_info_reads_dependencies_ip_slc_and_units(tmp_path):
    root = tmp_path / "ProjectRoot"
    for relative in ["unitlib", "projectlib", "nnelib", "SL_Library", "Configuration"]:
        (root / relative).mkdir(parents=True)

    z_file = root / "unitlib" / "PBSLC123.z"
    _write_text(z_file, "CommonLib.z\nIOHelpers.z\n")
    _write_text(z_file.with_suffix(".q"), '( Name "10.0.0.7" )\n')
    _write_text(z_file.with_suffix(".x"), 'pMixer : pType;\npFilter "UF" : pType ;\n')

    extractor = configgen.SattLineConfigExtractor(root)

    component = extractor.get_component_info(z_file, "Program", True, {"PBSLC123": "10.0.0.7"})

    assert component.name == "PBSLC123"
    assert component.type == "Program"
    assert component.dependencies == ["CommonLib", "IOHelpers"]
    assert component.ip_address == "10.0.0.7"
    assert component.slc == "SLC123"
    assert component.units == "(2) Filter, Mixer"


def test_configgen_workstation_mapper_is_case_insensitive():
    mapper = configgen.WorkstationMapper()

    assert mapper.get_workstations("kagc_op_utilf.z") == ["LOP01", "OP06", "OP07"]
    assert mapper.get_workstations("missing-station") == []
    assert mapper.get_physical_location("OPC01") == "Server Room"
    assert mapper.get_physical_location("UNKNOWN") == "Unknown Location"


def test_configgen_extractor_raises_when_required_directories_are_missing(tmp_path):
    with pytest.raises(ValueError, match="Required directories not found"):
        configgen.SattLineConfigExtractor(tmp_path / "missing-root")


def test_configgen_parse_all_configuration_files_handles_missing_and_empty_directory(tmp_path):
    root = tmp_path / "ProjectRoot"
    for relative in ["unitlib", "projectlib", "nnelib", "SL_Library"]:
        (root / relative).mkdir(parents=True)

    extractor = configgen.SattLineConfigExtractor(root)

    assert extractor.parse_all_configuration_files() == []

    extractor.kfiles_dir.mkdir()

    assert extractor.parse_all_configuration_files() == []


def test_configgen_parse_all_configuration_files_returns_only_valid_configs(tmp_path):
    root = tmp_path / "ProjectRoot"
    for relative in ["unitlib", "projectlib", "nnelib", "SL_Library", "Configuration"]:
        (root / relative).mkdir(parents=True)

    _write_text(
        root / "Configuration" / "Valid.k",
        'Configuration ( Version "v2" Date "2026-04-30" Name "Valid" )\n'
        'Program ( Name "Main.z" Directory "UnitLib" MainProgram True )\n',
    )
    _write_text(
        root / "Configuration" / "Invalid.k", 'Program ( Name "Broken.z" Directory "UnitLib" MainProgram True )\n'
    )

    extractor = configgen.SattLineConfigExtractor(root)

    result = extractor.parse_all_configuration_files()

    assert [config.config_name for config in result] == ["Valid"]


def test_configgen_get_z_files_returns_sorted_files_and_empty_for_missing_directory(tmp_path):
    root = tmp_path / "ProjectRoot"
    for relative in ["unitlib", "projectlib", "nnelib", "SL_Library", "Configuration"]:
        (root / relative).mkdir(parents=True)

    extractor = configgen.SattLineConfigExtractor(root)
    assert extractor.get_z_files(root / "does-not-exist") == []

    _write_text(root / "projectlib" / "B.z", "")
    _write_text(root / "projectlib" / "A.z", "")

    assert [path.name for path in extractor.get_z_files(root / "projectlib")] == ["A.z", "B.z"]


def test_configgen_read_dependencies_returns_empty_list_when_cached_read_fails(tmp_path, monkeypatch):
    root = tmp_path / "ProjectRoot"
    for relative in ["unitlib", "projectlib", "nnelib", "SL_Library", "Configuration"]:
        (root / relative).mkdir(parents=True)

    extractor = configgen.SattLineConfigExtractor(root)
    z_file = root / "unitlib" / "Broken.z"
    _write_text(z_file, "LibA.z\n")

    def _boom(_path):
        raise OSError("read failed")

    monkeypatch.setattr(extractor, "_read_file_cached", _boom)

    assert extractor.read_dependencies(z_file) == []


def test_configgen_ip_and_slc_helpers_cover_default_and_error_paths(tmp_path, monkeypatch):
    root = tmp_path / "ProjectRoot"
    for relative in ["unitlib", "projectlib", "nnelib", "SL_Library", "Configuration"]:
        (root / relative).mkdir(parents=True)

    z_file = root / "unitlib" / "WdSlc045.z"
    _write_text(z_file, "")
    extractor = configgen.SattLineConfigExtractor(root)

    assert extractor.get_ip_address(z_file) == "No Q-File"

    q_file = z_file.with_suffix(".q")
    _write_text(q_file, "No matching name token\n")
    assert extractor.get_ip_address(z_file) == "No SLC assigned"

    monkeypatch.setattr(configgen, "read_text_with_fallback", lambda _path: (_ for _ in ()).throw(OSError("q broken")))
    assert extractor.get_ip_address(z_file) == "Error reading Q-File"

    assert extractor.get_slc_name("No Q-File", {"PBSLC123": "10.0.0.7"}) == "No SLC"
    assert extractor.get_slc_name("10.0.0.8", {"WdSlc045": "10.0.0.8"}) == "SLC045"
    assert extractor.get_slc_name("10.0.0.9", {"OtherProgram": "10.0.0.9"}) == "No SLC"


def test_configgen_get_units_in_program_covers_missing_empty_and_error_paths(tmp_path, monkeypatch):
    root = tmp_path / "ProjectRoot"
    for relative in ["unitlib", "projectlib", "nnelib", "SL_Library", "Configuration"]:
        (root / relative).mkdir(parents=True)

    z_file = root / "unitlib" / "ProgramA.z"
    _write_text(z_file, "")
    extractor = configgen.SattLineConfigExtractor(root)

    assert extractor.get_units_in_program(z_file) == "No X-File"

    x_file = z_file.with_suffix(".x")
    _write_text(x_file, "No pType declarations here\n")
    assert extractor.get_units_in_program(z_file) == "No units assigned"

    monkeypatch.setattr(configgen, "read_text_with_fallback", lambda _path: (_ for _ in ()).throw(OSError("x broken")))
    assert extractor.get_units_in_program(z_file) == "Error reading X-File"


def test_configgen_get_component_info_sets_na_fields_for_non_program_components(tmp_path):
    root = tmp_path / "ProjectRoot"
    for relative in ["unitlib", "projectlib", "nnelib", "SL_Library", "Configuration"]:
        (root / relative).mkdir(parents=True)

    z_file = root / "projectlib" / "CommonLib.z"
    _write_text(z_file, "DepA.z\n")
    extractor = configgen.SattLineConfigExtractor(root)

    component = extractor.get_component_info(z_file, "Project Library", False, {})

    assert component.name == "CommonLib"
    assert component.ip_address == "N/A"
    assert component.slc == "N/A"
    assert component.units == "N/A"
    assert component.dependencies == ["DepA"]


def test_configgen_collect_and_process_components_cover_section_routing(tmp_path):
    class StubExtractor:
        unitlib_dir = Path("unitlib")
        projectlib_dir = Path("projectlib")
        nnelib_dir = Path("nnelib")
        sglib_dir = Path("SL_Library")

        def get_z_files(self, directory):
            mapping = {
                self.unitlib_dir: [Path("PBSLC100.z"), Path("IgnoreMe.z")],
                self.projectlib_dir: [Path("ProjectA.z")],
                self.nnelib_dir: [],
                self.sglib_dir: [Path("Support.z")],
            }
            return mapping[directory]

        def get_ip_address(self, z_file):
            return {"PBSLC100.z": "10.0.0.1", "IgnoreMe.z": "No Q-File"}[z_file.name]

        def get_component_info(self, z_file, component_type, has_ip, slc_programs):
            return configgen.ComponentInfo(
                name=z_file.stem,
                type=component_type,
                ip_address=slc_programs.get("PBSLC100", "N/A") if has_ip else "N/A",
                slc="SLC100" if z_file.stem == "PBSLC100" else "N/A",
                units="(1) UnitA" if has_ip else "N/A",
                dependencies=["Dep1"] if z_file.stem != "Support" else [],
            )

    generator = configgen.ExcelGenerator(extractor=_typed_extractor(StubExtractor()))

    slc_programs = generator._collect_slc_programs()
    component_data, dependency_data = generator._process_all_components(slc_programs)

    assert slc_programs == {"PBSLC100": "10.0.0.1"}
    assert [(component.name, component.type) for component in component_data] == [
        ("PBSLC100", "Program"),
        ("IgnoreMe", "Program"),
        ("ProjectA", "Project Library"),
        ("Support", "SG Library"),
    ]
    assert dependency_data == [("PBSLC100", "Dep1"), ("IgnoreMe", "Dep1"), ("ProjectA", "Dep1")]


def test_configgen_generate_runs_full_workbook_pipeline(tmp_path):
    output_path = tmp_path / "config.xlsx"

    class StubExtractor:
        unitlib_dir = Path("unitlib")
        projectlib_dir = Path("projectlib")
        nnelib_dir = Path("nnelib")
        sglib_dir = Path("SL_Library")

        def get_z_files(self, directory):
            mapping = {
                self.unitlib_dir: [Path("PBSLC100.z")],
                self.projectlib_dir: [Path("CommonLib.z")],
                self.nnelib_dir: [],
                self.sglib_dir: [],
            }
            return mapping[directory]

        def get_ip_address(self, z_file):
            return "10.0.0.1" if z_file.stem == "PBSLC100" else "No Q-File"

        def get_component_info(self, z_file, component_type, has_ip, slc_programs):
            return configgen.ComponentInfo(
                name=z_file.stem,
                type=component_type,
                ip_address=slc_programs.get(z_file.stem, "N/A") if has_ip else "N/A",
                slc="SLC100" if z_file.stem == "PBSLC100" else "N/A",
                units="(1) UnitA" if has_ip else "N/A",
                dependencies=["CommonLib"] if z_file.stem == "PBSLC100" else [],
            )

        def parse_all_configuration_files(self):
            return [
                configgen.ConfigurationFileInfo(
                    config_name="KaGC_Allf",
                    version="1.0",
                    date="2026-04-30",
                    main_program="PBSLC100",
                    programs=[{"name": "PBSLC100", "directory": "unitlib", "main_program": True}],
                    libraries=[{"name": "CommonLib", "directory": "projectlib"}],
                )
            ]

    generator = configgen.ExcelGenerator(extractor=_typed_extractor(StubExtractor()))
    generator.workstation_mapper.workstation_map = {"KaGC_Allf": ["OP11"]}
    generator.workstation_mapper.physical_locations = {"OP11": "Control Room 3"}

    generator.generate(output_path)

    assert output_path.exists()
    assert {sheet.title for sheet in generator.wb.worksheets} == {
        "Query Tool",
        "Dashboard",
        "System Components",
        "Library Dependencies",
        "Station Configuration",
        "Configuration Summary",
        "Configuration Details",
    }
    assert generator.components_ws["B2"].value == "PBSLC100"
    assert generator.dependencies_ws["D2"].value == "CommonLib"


def test_configgen_main_returns_2_for_invalid_root(tmp_path, capsys):
    invalid_root = tmp_path / "missing-root"

    result = configgen.main([str(invalid_root)])

    captured = capsys.readouterr()
    assert result == 2
    assert "Invalid root directory" in captured.err


def test_configgen_main_generates_output_with_stubbed_components(tmp_path, monkeypatch, capsys):
    root_dir = tmp_path / "root"
    root_dir.mkdir()
    output_path = tmp_path / "out.xlsx"
    captured = {}

    class StubExtractor:
        def __init__(self, root):
            captured["extractor_root"] = root

    class StubGenerator:
        def __init__(self, extractor):
            captured["generator_extractor"] = extractor

        def generate(self, output):
            captured["output"] = output

    monkeypatch.setattr(configgen, "SattLineConfigExtractor", StubExtractor)
    monkeypatch.setattr(configgen, "ExcelGenerator", StubGenerator)

    result = configgen.main([str(root_dir), "--output", str(output_path)])

    captured_output = capsys.readouterr().out
    assert result == 0
    assert captured["extractor_root"] == root_dir.resolve()
    assert captured["output"] == output_path.resolve()
    assert "Configuration Excel file generated successfully" in captured_output


def test_configgen_style_manager_applies_header_styling():
    from openpyxl import Workbook

    wb = Workbook()
    ws = _active_worksheet(wb)

    style_manager = configgen.StyleManager()
    cell = ws["A1"]
    cell.value = "Test Header"
    style_manager.apply_header_style(cell)

    assert cell.font.bold is True
    assert cell.font.size == 11
    assert "4472C4" in cell.fill.start_color.rgb


def test_configgen_worksheet_helper_setup_headers():
    from openpyxl import Workbook

    wb = Workbook()
    ws = _active_worksheet(wb)
    style_manager = configgen.StyleManager()
    headers = ["ID", "Name", "Value"]

    configgen.WorksheetHelper.setup_headers(ws, headers, style_manager)

    assert ws["A1"].value == "ID"
    assert ws["B1"].value == "Name"
    assert ws["C1"].value == "Value"
    assert ws["A1"].font.bold is True


def test_configgen_worksheet_helper_create_table():
    from openpyxl import Workbook

    wb = Workbook()
    ws = _active_worksheet(wb)

    configgen.WorksheetHelper.create_table(ws, "TestTable", "A1:C3")

    assert len(ws.tables) == 1
    assert ws.tables["TestTable"] is not None


def test_configgen_worksheet_helper_auto_fit_columns():
    from openpyxl import Workbook

    wb = Workbook()
    ws = _active_worksheet(wb)
    ws["A1"] = "Short"
    ws["B1"] = "This is a much longer value that should increase column width"

    configgen.WorksheetHelper.auto_fit_columns(ws, max_width=60)

    assert ws.column_dimensions["A"].width > 0
    assert ws.column_dimensions["B"].width >= ws.column_dimensions["A"].width


def test_configgen_aggregate_slc_numbers_matches_component_case_insensitive():
    component_data = [
        configgen.ComponentInfo(
            name="PBSLC123", type="Program", ip_address="10.0.0.1", slc="SLC123", units="(2) A, B", dependencies=[]
        ),
        configgen.ComponentInfo(
            name="OpsStation", type="Program", ip_address="10.0.0.2", slc="No SLC", units="(0) ", dependencies=[]
        ),
    ]
    config = {"programs": ["pbslc123", "OpsStation"]}
    generator = configgen.ExcelGenerator.__new__(configgen.ExcelGenerator)

    result = generator._aggregate_slc_numbers(config, component_data)

    assert result == "SLC123"


def test_configgen_aggregate_units_extracts_from_format():
    component_data = [
        configgen.ComponentInfo(
            name="Prog1",
            type="Program",
            ip_address="10.0.0.1",
            slc="SLC123",
            units="(3) UnitA, UnitB, UnitC",
            dependencies=[],
        ),
        configgen.ComponentInfo(
            name="Prog2",
            type="Program",
            ip_address="10.0.0.2",
            slc="No SLC",
            units="(1) UnitD",
            dependencies=[],
        ),
    ]
    config = {"programs": ["Prog1", "Prog2"]}
    generator = configgen.ExcelGenerator.__new__(configgen.ExcelGenerator)

    result = generator._aggregate_units(config, component_data)

    assert "UnitA" in result
    assert "UnitB" in result
    assert "UnitC" in result
    assert "UnitD" in result
    assert result.startswith("(")


def test_configgen_format_units_for_station_returns_no_units_for_invalid_input():
    generator = configgen.ExcelGenerator.__new__(configgen.ExcelGenerator)

    assert generator._format_units_for_station("") == "No units assigned"
    assert generator._format_units_for_station("No X-File") == "No units assigned"
    assert generator._format_units_for_station("Error reading X-File") == "No units assigned"


def test_configgen_format_units_for_station_extracts_from_paren_format():
    generator = configgen.ExcelGenerator.__new__(configgen.ExcelGenerator)

    result = generator._format_units_for_station("(3) UnitA, UnitB, UnitC")

    assert result == "UnitA, UnitB, UnitC"


def test_configgen_determine_station_type_classifies_by_prefix():
    generator = configgen.ExcelGenerator.__new__(configgen.ExcelGenerator)

    assert generator._determine_station_type("LOP01") == "Local Operator Panel"
    assert generator._determine_station_type("OP06") == "Operator Station"
    assert generator._determine_station_type("OPC01") == "Operator Station"
    assert generator._determine_station_type("PRG01") == "Programmer Station"
    assert generator._determine_station_type("Journal Server 1 Primary") == "Journal Server"
    assert generator._determine_station_type("UNKNOWN") == "Special System"


def test_configgen_populate_components_sheet_writes_rows_and_returns_count():
    generator = configgen.ExcelGenerator(extractor=_typed_extractor(object()))
    component_data = [
        configgen.ComponentInfo(
            name="PBSLC123",
            type="Program",
            ip_address="10.0.0.1",
            slc="SLC123",
            units="(2) UnitA, UnitB",
            dependencies=["CommonLib"],
        ),
        configgen.ComponentInfo(
            name="CommonLib",
            type="Project Library",
            ip_address="N/A",
            slc="No SLC",
            units="No X-File",
            dependencies=[],
        ),
    ]

    row_count = generator._populate_components_sheet(component_data)

    assert row_count == 3
    assert generator.components_ws["A1"].value == "ID"
    assert generator.components_ws["B2"].value == "PBSLC123"
    assert generator.components_ws["F3"].value == "Project Library"


def test_configgen_populate_dependencies_sheet_uses_case_insensitive_type_lookup():
    generator = configgen.ExcelGenerator(extractor=_typed_extractor(object()))
    generator.all_component_data = [
        configgen.ComponentInfo(
            name="PBSLC123",
            type="Program",
            ip_address="10.0.0.1",
            slc="SLC123",
            units="(1) UnitA",
            dependencies=["CommonLib"],
        ),
        configgen.ComponentInfo(
            name="CommonLib",
            type="Project Library",
            ip_address="N/A",
            slc="No SLC",
            units="No X-File",
            dependencies=[],
        ),
    ]

    row_count = generator._populate_dependencies_sheet([("pbslc123", "commonlib"), ("UnknownComp", "UnknownLib")])

    assert row_count == 3
    assert generator.dependencies_ws["C2"].value == "Program"
    assert generator.dependencies_ws["E2"].value == "Project Library"
    assert generator.dependencies_ws["C3"].value == "Unknown"
    assert generator.dependencies_ws["E3"].value == "Unknown"


def test_configgen_create_dashboard_sets_title_and_merge():
    generator = configgen.ExcelGenerator(extractor=_typed_extractor(object()))

    generator._create_dashboard(2, 2)

    assert "A1:H1" in [str(rng) for rng in generator.dashboard_ws.merged_cells.ranges]
    assert generator.dashboard_ws["A1"].value == "SattLine Configuration Dashboard"
    assert generator.dashboard_ws.row_dimensions[1].height == 30


def test_configgen_create_configuration_summary_sheet_no_configs_writes_message():
    class StubExtractor:
        def parse_all_configuration_files(self):
            return []

    generator = configgen.ExcelGenerator(extractor=_typed_extractor(StubExtractor()))

    generator._create_configuration_summary_sheet()

    ws = generator.wb["Configuration Summary"]
    assert ws["A3"].value == "No configuration files found or parsing failed"


def test_configgen_create_configuration_summary_sheet_builds_rows_and_table():
    class StubExtractor:
        def parse_all_configuration_files(self):
            return [
                configgen.ConfigurationFileInfo(
                    config_name="kBConfig.z",
                    version="1.0",
                    date="2026-04-29",
                    main_program="ProgB",
                    programs=[{"name": "ProgB", "directory": "unitlib", "main_program": True}],
                    libraries=[{"name": "LibB", "directory": "projectlib"}],
                ),
                configgen.ConfigurationFileInfo(
                    config_name="kAConfig.z",
                    version="2.0",
                    date="2026-04-30",
                    main_program="ProgA",
                    programs=[{"name": "ProgA", "directory": "unitlib", "main_program": True}],
                    libraries=[{"name": "LibA", "directory": "projectlib"}],
                ),
            ]

    generator = configgen.ExcelGenerator(extractor=_typed_extractor(StubExtractor()))

    generator._create_configuration_summary_sheet()

    ws = generator.wb["Configuration Summary"]
    assert ws["A4"].value == "kAConfig.z"
    assert ws["A5"].value == "kBConfig.z"
    assert "ConfigurationSummary" in ws.tables


def test_configgen_create_configuration_details_sheet_builds_program_and_library_rows():
    class StubExtractor:
        def parse_all_configuration_files(self):
            return [
                configgen.ConfigurationFileInfo(
                    config_name="kMain.z",
                    version="3.1",
                    date="2026-04-29",
                    main_program="ProgMain",
                    programs=[{"name": "ProgMain", "directory": "unitlib", "main_program": True}],
                    libraries=[{"name": "CommonLib", "directory": "projectlib"}],
                )
            ]

    generator = configgen.ExcelGenerator(extractor=_typed_extractor(StubExtractor()))

    generator._create_configuration_details_sheet()

    ws = generator.wb["Configuration Details"]
    assert ws["A4"].value == "kMain.z"
    assert ws["B4"].value == "Program"
    assert ws["B5"].value == "Library"
    assert "ConfigurationDetails" in ws.tables


def test_configgen_build_station_configurations_adds_transitive_non_program_libraries():
    class StubExtractor:
        def parse_all_configuration_files(self):
            return [
                configgen.ConfigurationFileInfo(
                    config_name="KaGc_AllF",
                    version="1.0",
                    date="2026-04-29",
                    main_program="ProgA",
                    programs=[{"name": "ProgA", "directory": "unitlib", "main_program": True}],
                    libraries=[{"name": "DeclaredLib", "directory": "projectlib"}],
                )
            ]

    generator = configgen.ExcelGenerator(extractor=_typed_extractor(StubExtractor()))
    generator.workstation_mapper.workstation_map = {"KaGC_Allf": ["OP11"]}
    generator.workstation_mapper.physical_locations = {"OP11": "Control Room 3"}

    component_data = [
        configgen.ComponentInfo(
            name="proga",
            type="Program",
            ip_address="10.0.0.1",
            slc="SLC100",
            units="(1) UnitA",
            dependencies=["declaredlib", "TransitiveLib", "ProgB"],
        ),
        configgen.ComponentInfo(
            name="DeclaredLib",
            type="Project Library",
            ip_address="N/A",
            slc="No SLC",
            units="N/A",
            dependencies=[],
        ),
        configgen.ComponentInfo(
            name="transitivelib",
            type="SG Library",
            ip_address="N/A",
            slc="No SLC",
            units="N/A",
            dependencies=[],
        ),
        configgen.ComponentInfo(
            name="ProgB",
            type="Program",
            ip_address="10.0.0.2",
            slc="SLC101",
            units="(1) UnitB",
            dependencies=[],
        ),
    ]

    result = generator._build_station_configurations(component_data)

    assert result["OP11"]["config_file"] == "KaGC_Allf"
    assert result["OP11"]["type"] == "Operator Station"
    assert result["OP11"]["programs"] == ["ProgA"]
    assert result["OP11"]["libraries"] == ["DeclaredLib", "TransitiveLib"]


def test_configgen_create_station_configuration_sheet_writes_defaults_for_missing_config_data():
    class StubExtractor:
        def parse_all_configuration_files(self):
            return []

    generator = configgen.ExcelGenerator(extractor=_typed_extractor(StubExtractor()))
    generator.workstation_mapper.workstation_map = {"KaGC_Missing": ["OP99"]}
    generator.workstation_mapper.physical_locations = {"OP99": "Lab"}

    generator._create_station_configuration_sheet([])

    ws = generator.wb["Station Configuration"]
    assert ws["A4"].value == "OP99"
    assert ws["B4"].value == "Operator Station"
    assert ws["C4"].value == "Lab"
    assert ws["D4"].value == "KaGC_Missing"
    assert ws["E4"].value == "No SLC"
    assert ws["F4"].value == "No programs"
    assert ws["G4"].value == "No libraries"
    assert ws["H4"].value == "No units assigned"
    assert ws["I4"].value == "N/A"
    assert "StationConfiguration" in ws.tables


def test_get_example_fixtures_for_analyzer_returns_examples():
    from sattlint.docgenerator.analyzer_ref import get_example_fixtures_for_analyzer

    variables_examples = get_example_fixtures_for_analyzer("variables")
    assert len(variables_examples) >= 1
    assert any("CommonQualityIssues" in ex["fixture"] for ex in variables_examples)

    sfc_examples = get_example_fixtures_for_analyzer("sfc")
    assert len(sfc_examples) >= 2
    assert any("ParallelWriteRace" in ex["fixture"] for ex in sfc_examples)

    unknown_examples = get_example_fixtures_for_analyzer("nonexistent")
    assert unknown_examples == []


def test_get_example_fixtures_for_analyzer_returns_expected_rule_ids():
    from sattlint.docgenerator.analyzer_ref import get_example_fixtures_for_analyzer

    shadowing_examples = get_example_fixtures_for_analyzer("shadowing")
    assert len(shadowing_examples) >= 1
    example = shadowing_examples[0]
    assert "semantic.shadowing" in example["expected_rule_ids"]


def test_build_analyzer_reference_entry():
    from sattlint.analyzers.framework import AnalyzerSpec
    from sattlint.docgenerator.analyzer_ref import build_analyzer_reference_entry

    def _dummy_run(context):
        from sattlint.analyzers.framework import SimpleReport

        return SimpleReport(name="test")

    spec = AnalyzerSpec(
        key="variables",
        name="Variable issues",
        description="Unused/read-only/never-read variables",
        run=_dummy_run,
        enabled=True,
        supports_live_diagnostics=True,
    )

    entry = build_analyzer_reference_entry(spec)

    assert entry["key"] == "variables"
    assert entry["name"] == "Variable issues"
    assert entry["enabled"] is True
    assert "delivery" in entry
    assert "rules" in entry
    assert "examples" in entry
    assert entry["example_count"] >= 1


def test_build_full_analyzer_reference():
    from sattlint.docgenerator.analyzer_ref import build_full_analyzer_reference

    reference = build_full_analyzer_reference()

    assert "generated_by" in reference
    assert "schema_version" in reference
    assert "analyzers" in reference
    assert "total_analyzers" in reference
    assert "total_rules" in reference
    assert reference["total_analyzers"] > 0
    assert reference["total_rules"] > 0


def test_render_analyzer_reference_markdown():
    from sattlint.docgenerator.analyzer_ref import render_analyzer_reference_markdown

    markdown = render_analyzer_reference_markdown()

    assert "# SattLint Analyzer Reference" in markdown
    assert "Generated by:" in markdown
    assert "Total analyzers:" in markdown
    assert "Total rules:" in markdown


def test_render_analyzer_reference_markdown_contains_analyzer_info():
    from sattlint.docgenerator.analyzer_ref import render_analyzer_reference_markdown

    markdown = render_analyzer_reference_markdown()

    assert "## Variable issues" in markdown or "## SFC checks" in markdown
    assert "**Description:**" in markdown
    assert "**Enabled:**" in markdown


def test_save_analyzer_reference_json(tmp_path):
    from sattlint.docgenerator.analyzer_ref import save_analyzer_reference_json

    output_path = tmp_path / "analyzer_ref.json"
    save_analyzer_reference_json(output_path)

    assert output_path.exists()
    import json

    data = json.loads(output_path.read_text())
    assert "analyzers" in data
    assert len(data["analyzers"]) > 0


def test_save_analyzer_reference_markdown(tmp_path):
    from sattlint.docgenerator.analyzer_ref import save_analyzer_reference_markdown

    output_path = tmp_path / "analyzer_ref.md"
    save_analyzer_reference_markdown(output_path)

    assert output_path.exists()
    content = output_path.read_text()
    assert "# SattLint Analyzer Reference" in content
    assert "Generated by:" in content

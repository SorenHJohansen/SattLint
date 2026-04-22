"""Generate a Functional-Specification-style Word document from a SattLine AST."""

from __future__ import annotations

import logging
import pathlib
import re
import typing as t
from dataclasses import dataclass

from docx import Document as DocumentFactory
from docx.document import Document as DocClass
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from sattline_parser.utils.formatter import format_expr

from .. import config as config_module
from ..analyzers.framework import Issue
from ..models.ast_model import (
    Sequence,
    SFCAlternative,
    SFCBreak,
    SFCFork,
    SFCParallel,
    SFCStep,
    SFCSubsequence,
    SFCTransition,
    SFCTransitionSub,
    Variable,
)
from .classification import (
    DocumentationClassification,
    DocumentedModule,
    classify_documentation_structure,
    discover_documentation_unit_candidates,
)

log = logging.getLogger("SattLint")

_TAG_NAME_RE = re.compile(r"^[A-Z]{1,4}\d{3,4}$")
_DEVICE_PREFIX_RE = re.compile(r"^(?:V|TT|PT|FT|WT|AIT|LS|ZS|PC|LC|TC|PM|BM)\d{3,4}$", re.IGNORECASE)
_OPERATION_EXCLUDE_NAMES = {"mes_info", "mes_stop"}


@dataclass(frozen=True)
class DocumentUnit:
    root: DocumentedModule
    unit_code: str
    title: str
    unit_class: str
    section_name: str
    equipment_modules: list[DocumentedModule]
    operations: list[DocumentedModule]
    recipe_parameters: list[DocumentedModule]
    engineering_parameters: list[DocumentedModule]
    user_parameters: list[DocumentedModule]


@dataclass(frozen=True)
class SequenceRenderRow:
    node_type: str
    name: str
    detail: str
    enter: str = ""
    active: str = ""
    exit: str = ""


def _ensure_styles(doc: DocClass) -> None:
    styles = doc.styles
    if "Underlined" not in styles:
        style = t.cast(t.Any, styles.add_style("Underlined", WD_STYLE_TYPE.PARAGRAPH))
        style.base_style = styles["Normal"]
        style.font.underline = True
        style.font.size = Pt(10)


def _heading(doc: DocClass, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _paragraph(
    doc: DocClass,
    text: str,
    *,
    bold: bool = False,
    style: str | None = None,
) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.bold = bold


def _centered_title(doc: DocClass, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def _bullet(doc: DocClass, text: str) -> None:
    doc.add_paragraph(text, style="List Paragraph")


def _underlined(doc: DocClass, text: str) -> None:
    _paragraph(doc, text, style="Underlined")


def _table(
    doc: DocClass,
    headers: list[str],
    rows: t.Sequence[t.Sequence[object]],
    col_widths: tuple[int, ...] = (),
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)

    if col_widths:
        for index, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[index].width = Inches(width)


def _value_text(value: object | None) -> str:
    if value is None:
        return ""
    return str(value)


def _format_coord(coord: tuple[float, ...] | None) -> str:
    if coord is None:
        return "<none>"
    return "(" + ", ".join(f"{value:.6g}" for value in coord) + ")"


def _prettify_name(text: str) -> str:
    if not text:
        return ""
    normalized = text.replace("_", " ").strip()
    if normalized.isupper():
        return normalized
    return normalized


def _sequence_code_text(statements: list[object]) -> str:
    return "\n".join(format_expr(statement) for statement in statements)


def _append_sequence_rows(
    rows: list[SequenceRenderRow],
    nodes: list[object],
    *,
    context: str | None = None,
) -> None:
    for node in nodes:
        detail_prefix = f"{context}: " if context else ""
        if isinstance(node, SFCStep):
            rows.append(
                SequenceRenderRow(
                    node_type="Init step" if node.kind == "init" else "Step",
                    name=node.name,
                    detail=detail_prefix + ("Initial step" if node.kind == "init" else "Execution step"),
                    enter=_sequence_code_text(node.code.enter),
                    active=_sequence_code_text(node.code.active),
                    exit=_sequence_code_text(node.code.exit),
                )
            )
            continue

        if isinstance(node, SFCTransition):
            rows.append(
                SequenceRenderRow(
                    node_type="Transition",
                    name=node.name or "<unnamed>",
                    detail=detail_prefix + format_expr(node.condition),
                )
            )
            continue

        if isinstance(node, SFCAlternative):
            rows.append(
                SequenceRenderRow(
                    node_type="Alternative",
                    name="",
                    detail=detail_prefix + f"{len(node.branches)} branches",
                )
            )
            for index, branch in enumerate(node.branches, start=1):
                branch_context = f"Alternative branch {index}"
                rows.append(
                    SequenceRenderRow(
                        node_type="Branch",
                        name=f"Branch {index}",
                        detail=detail_prefix + "Alternative path",
                    )
                )
                _append_sequence_rows(rows, branch, context=branch_context)
            continue

        if isinstance(node, SFCParallel):
            rows.append(
                SequenceRenderRow(
                    node_type="Parallel",
                    name="",
                    detail=detail_prefix + f"{len(node.branches)} branches",
                )
            )
            for index, branch in enumerate(node.branches, start=1):
                branch_context = f"Parallel branch {index}"
                rows.append(
                    SequenceRenderRow(
                        node_type="Branch",
                        name=f"Branch {index}",
                        detail=detail_prefix + "Parallel path",
                    )
                )
                _append_sequence_rows(rows, branch, context=branch_context)
            continue

        if isinstance(node, SFCSubsequence):
            rows.append(
                SequenceRenderRow(
                    node_type="Subsequence",
                    name=node.name,
                    detail=detail_prefix + "Nested sequence",
                )
            )
            _append_sequence_rows(rows, node.body, context=f"Subsequence {node.name}")
            continue

        if isinstance(node, SFCTransitionSub):
            rows.append(
                SequenceRenderRow(
                    node_type="Transition section",
                    name=node.name,
                    detail=detail_prefix + "Nested transition sequence",
                )
            )
            _append_sequence_rows(rows, node.body, context=f"Transition section {node.name}")
            continue

        if isinstance(node, SFCFork):
            rows.append(
                SequenceRenderRow(
                    node_type="Fork",
                    name=node.target,
                    detail=detail_prefix + "Fork target",
                )
            )
            continue

        if isinstance(node, SFCBreak):
            rows.append(
                SequenceRenderRow(
                    node_type="Break",
                    name="",
                    detail=detail_prefix + "Break sequence flow",
                )
            )
            continue

        rows.append(
            SequenceRenderRow(
                node_type="Statement",
                name="",
                detail=detail_prefix + format_expr(node),
            )
        )


def _sequence_render_rows(sequence: Sequence) -> list[SequenceRenderRow]:
    rows: list[SequenceRenderRow] = []
    _append_sequence_rows(rows, list(sequence.code or []))
    return rows


def _sequence_table_rows(rows: list[SequenceRenderRow]) -> list[list[str]]:
    return [[row.node_type, row.name, row.detail, row.enter, row.active, row.exit] for row in rows]


def _mapping_target_name(mapping) -> str:
    if isinstance(mapping.target, dict):
        return str(mapping.target.get("var_name", ""))
    return str(mapping.target)


def _mapping_source_text(mapping) -> str:
    if isinstance(mapping.source, dict):
        return str(mapping.source.get("var_name", ""))
    if mapping.source is not None:
        return str(mapping.source)
    if mapping.source_literal is None:
        return ""
    return str(mapping.source_literal)


def _mapping_value(entry: DocumentedModule, *target_names: str) -> str | None:
    for wanted_name in target_names:
        wanted_cf = wanted_name.casefold()
        for mapping in entry.parametermappings:
            target_name = _mapping_target_name(mapping).casefold()
            if target_name != wanted_cf:
                continue
            source_text = _mapping_source_text(mapping).strip()
            if source_text:
                return source_text
    return None


def _display_name(entry: DocumentedModule) -> str:
    return _mapping_value(entry, "HeaderName", "MediaName", "Name") or entry.name


def _module_title(entry: DocumentedModule) -> str:
    return _mapping_value(entry, "Name") or entry.name


def _module_description(entry: DocumentedModule) -> str:
    description = _mapping_value(entry, "Name")
    if description and description.casefold() != entry.name.casefold():
        return f"Configured as {description}."
    for variable in (*entry.moduleparameters, *entry.localvariables):
        if variable.description:
            return variable.description
    return f"Detected {_prettify_name(entry.moduletype_name or entry.kind)} instance."


def _moduletype_summary(entry: DocumentedModule) -> str:
    return entry.moduletype_label or entry.moduletype_name or entry.kind


def _build_units(classification: DocumentationClassification) -> list[DocumentUnit]:
    scope_roots = list(classification.scope.roots or []) if classification.scope else []
    roots = scope_roots or discover_documentation_unit_candidates(classification)
    units: list[DocumentUnit] = []
    for root in roots:
        units.append(
            DocumentUnit(
                root=root,
                unit_code=_mapping_value(root, "Name") or root.name,
                title=_display_name(root),
                unit_class=root.moduletype_name or root.moduletype_label or root.name,
                section_name=_mapping_value(root, "SectionName") or root.name,
                equipment_modules=_unit_category_descendants(root, classification, "em"),
                operations=_unit_category_descendants(root, classification, "ops"),
                recipe_parameters=_unit_category_descendants(root, classification, "rp", top_level_only=True),
                engineering_parameters=_unit_category_descendants(root, classification, "ep", top_level_only=True),
                user_parameters=_unit_category_descendants(root, classification, "up", top_level_only=True),
            )
        )
    return units


def _unit_category_descendants(
    root: DocumentedModule,
    classification: DocumentationClassification,
    category: str,
    *,
    top_level_only: bool = False,
) -> list[DocumentedModule]:
    entries = classification.descendants(root, category=category)
    if not top_level_only:
        return entries

    excluded_ancestors = classification.descendants(root, category="em") + classification.descendants(
        root, category="ops"
    )
    return [
        entry
        for entry in entries
        if not any(
            entry.path != ancestor.path and entry.path[: len(ancestor.path)] == ancestor.path
            for ancestor in excluded_ancestors
        )
    ]


def _descendants(root: DocumentedModule, classification: DocumentationClassification) -> list[DocumentedModule]:
    return classification.descendants(root)


def _support_entries(unit: DocumentUnit, classification: DocumentationClassification) -> list[DocumentedModule]:
    descendants = _descendants(unit.root, classification)
    excluded = [*unit.equipment_modules, *unit.operations]
    return [
        entry
        for entry in descendants
        if not any(
            entry.path != ancestor.path and entry.path[: len(ancestor.path)] == ancestor.path for ancestor in excluded
        )
    ]


def _parameter_catalog_rows(entries: list[DocumentedModule]) -> list[list[str]]:
    rows: list[list[str]] = []
    for entry in entries:
        rows.append(
            [
                _module_title(entry),
                _moduletype_summary(entry),
                entry.short_path,
                _module_description(entry),
            ]
        )
    return rows


def _module_list_rows(entries: list[DocumentedModule]) -> list[list[str]]:
    rows: list[list[str]] = []
    for entry in entries:
        rows.append(
            [
                _module_title(entry),
                _module_description(entry),
                f"See section {_module_title(entry)}",
            ]
        )
    return rows


def _metadata_value(entry: DocumentedModule, *target_names: str) -> str:
    return _mapping_value(entry, *target_names) or ""


def _first_non_empty(*values: str) -> str:
    for value in values:
        if value.strip():
            return value
    return ""


def _entry_variable(entry: DocumentedModule, *names: str) -> Variable | None:
    wanted = {name.casefold() for name in names}
    for variable in (*entry.moduleparameters, *entry.localvariables):
        if variable.name.casefold() in wanted:
            return variable
    return None


def _entry_variable_text(entry: DocumentedModule, *names: str) -> str:
    variable = _entry_variable(entry, *names)
    if variable is None:
        return ""
    return _first_non_empty(_value_text(variable.init_value), variable.description or "")


def _configurable_parameter_rows(unit: DocumentUnit) -> list[list[str]]:
    rows: list[list[str]] = []
    ignored = {
        "p",
        "allow",
        "colours",
        "programname",
        "headernamecolour",
        "nextview",
    }
    descriptions = {variable.name.casefold(): variable.description or "" for variable in unit.root.moduleparameters}
    for mapping in unit.root.parametermappings:
        target_name = _mapping_target_name(mapping).strip()
        if not target_name or target_name.casefold() in ignored:
            continue
        source_text = _mapping_source_text(mapping).strip()
        if not source_text:
            continue
        rows.append(
            [
                target_name,
                descriptions.get(target_name.casefold(), ""),
                unit.unit_code,
                source_text,
            ]
        )
    return rows


def _measurement_rows(entries: list[DocumentedModule]) -> list[list[str]]:
    rows: list[list[str]] = []
    for entry in entries:
        rows.append(
            [
                _module_title(entry),
                entry.name,
                _entry_variable_text(entry, "min", "lowlimit", "range_min"),
                _entry_variable_text(entry, "max", "highlimit", "range_max"),
                _first_non_empty(
                    _metadata_value(entry, "EngUnit", "Unit", "EU"),
                    _entry_variable_text(entry, "engunit", "unit", "eu"),
                ),
                _first_non_empty(
                    _metadata_value(entry, "LogIntervalMax", "MaxLogInterval"),
                    _entry_variable_text(entry, "logintervalmax", "maxloginterval"),
                ),
                _first_non_empty(
                    _metadata_value(entry, "DeadbandRelative", "DeadBand"),
                    _entry_variable_text(entry, "deadbandrelative", "deadband"),
                ),
                _first_non_empty(
                    _metadata_value(entry, "LogIntervalMin", "MinLogInterval"),
                    _entry_variable_text(entry, "logintervalmin", "minloginterval"),
                ),
            ]
        )
    return rows


def _generic_section_rows(entries: list[DocumentedModule]) -> list[list[str]]:
    return [
        [
            _module_title(entry),
            _moduletype_summary(entry),
            entry.short_path,
            _module_description(entry),
        ]
        for entry in entries
    ]


def _simple_name_tag_rows(entries: list[DocumentedModule], description_label: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for entry in entries:
        rows.append(
            [
                _module_description(entry) if description_label == "description" else _module_title(entry),
                entry.name,
            ]
        )
    return rows


def _event_table_rows(entries: list[DocumentedModule]) -> list[list[str]]:
    rows: list[list[str]] = []
    for entry in entries:
        rows.append(
            [
                entry.name,
                _module_description(entry),
                _metadata_value(entry, "Severity", "Sev"),
                _metadata_value(entry, "Activation", "Condition"),
            ]
        )
    return rows


def _interlock_rows(variables: list[Variable]) -> list[list[str]]:
    return [
        [
            variable.name,
            variable.description or "",
            _value_text(variable.init_value),
            "",
        ]
        for variable in variables
    ]


def _exception_rows(variables: list[Variable]) -> list[list[str]]:
    return [
        [
            variable.name,
            variable.description or "",
        ]
        for variable in variables
    ]


def _calculation_rows(entries: list[DocumentedModule]) -> list[list[str]]:
    rows: list[list[str]] = []
    for entry in entries:
        rows.append(
            [
                entry.name,
                _module_description(entry),
                entry.short_path,
            ]
        )
    return rows


def _communication_rows(unit: DocumentUnit, *, direction: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for mapping in unit.root.parametermappings:
        target_name = _mapping_target_name(mapping).strip()
        if not target_name:
            continue
        target_cf = target_name.casefold()
        if direction == "from" and not target_cf.startswith("inlet"):
            continue
        if direction == "to" and not target_cf.startswith("outlet"):
            continue
        connection = _mapping_source_text(mapping).strip() or target_name
        rows.append(
            [
                connection,
                f"Mapped via {target_name}",
            ]
        )
    return rows


def _state_rows(entry: DocumentedModule, classification: DocumentationClassification) -> list[list[str]]:
    rows: list[list[str]] = []
    seen_paths: set[tuple[str, ...]] = set()
    for descendant in classification.descendants(entry):
        if descendant.path in seen_paths:
            continue
        if len(descendant.path) < 2:
            continue
        if descendant.path[-2].casefold() != "panel":
            continue
        if descendant.name.casefold().startswith("kahctoggle"):
            continue
        rows.append(
            [
                _module_title(descendant),
                _module_description(descendant),
                _state_logic_summary(descendant, classification),
            ]
        )
        seen_paths.add(descendant.path)
    return rows


def _state_logic_summary(entry: DocumentedModule, classification: DocumentationClassification) -> str:
    sequence_count = 0
    equation_count = 0
    state_logic_types = 0
    for descendant in classification.descendants(entry):
        if descendant.moduletype_name and descendant.moduletype_name.casefold() == "statelogic":
            state_logic_types += 1
        if descendant.modulecode is not None:
            sequence_count += len(descendant.modulecode.sequences or [])
            equation_count += len(descendant.modulecode.equations or [])
    parts: list[str] = []
    if state_logic_types:
        parts.append(f"{state_logic_types} state-logic modules")
    if sequence_count:
        parts.append(f"{sequence_count} sequences")
    if equation_count:
        parts.append(f"{equation_count} equation blocks")
    return ", ".join(parts) if parts else "Display-only state definition"


def _pid_controller_rows(entry: DocumentedModule, classification: DocumentationClassification) -> list[list[str]]:
    rows: list[list[str]] = []
    for descendant in classification.descendants(entry):
        name_cf = descendant.name.casefold()
        label_cf = _moduletype_summary(descendant).casefold()
        if not ("pid" in label_cf or label_cf.endswith("ctrl") or name_cf.startswith(("pc", "lc", "tc"))):
            continue
        rows.append(
            [
                descendant.name,
                _moduletype_summary(descendant),
                descendant.short_path,
                _module_description(descendant),
            ]
        )
    return rows


def _sequence_rows(
    entry: DocumentedModule, classification: DocumentationClassification
) -> list[tuple[str, list[SequenceRenderRow]]]:
    rows: list[tuple[str, list[SequenceRenderRow]]] = []
    seen: set[tuple[str, str]] = set()
    for candidate in [entry, *classification.descendants(entry)]:
        if candidate.modulecode is None:
            continue
        for sequence in candidate.modulecode.sequences or []:
            key = (candidate.short_path, sequence.name)
            if key in seen:
                continue
            rows.append((sequence.name, _sequence_render_rows(sequence)))
            seen.add(key)
    return rows


def _message_rows(entry: DocumentedModule, classification: DocumentationClassification) -> list[list[str]]:
    rows: list[list[str]] = []
    for descendant in classification.descendants(entry):
        name_cf = descendant.name.casefold()
        label_cf = _moduletype_summary(descendant).casefold()
        if "message" not in name_cf and "message" not in label_cf and "opmess" not in label_cf:
            continue
        rows.append(
            [
                descendant.name,
                _moduletype_summary(descendant),
                descendant.short_path,
                _module_description(descendant),
            ]
        )
    return rows


def _event_rows(entry: DocumentedModule, classification: DocumentationClassification) -> list[list[str]]:
    rows: list[list[str]] = []
    for descendant in classification.descendants(entry):
        name_cf = descendant.name.casefold()
        label_cf = _moduletype_summary(descendant).casefold()
        if not (name_cf.startswith("event") or "journal" in label_cf or "event" in label_cf):
            continue
        rows.append(
            [
                descendant.name,
                _moduletype_summary(descendant),
                descendant.short_path,
                _module_description(descendant),
            ]
        )
    return rows


def _special_logging_rows(entry: DocumentedModule, classification: DocumentationClassification) -> list[list[str]]:
    rows: list[list[str]] = []
    for descendant in classification.descendants(entry):
        label_cf = _moduletype_summary(descendant).casefold()
        name_cf = descendant.name.casefold()
        if "journal" not in label_cf and "log" not in name_cf and "journal" not in name_cf:
            continue
        rows.append(
            [
                descendant.name,
                _moduletype_summary(descendant),
                descendant.short_path,
                _module_description(descendant),
            ]
        )
    return rows


def _render_named_table_section(
    doc: DocClass,
    title: str,
    rows: list[list[str]],
    *,
    headers: list[str],
    level: int = 3,
    widths: tuple[int, ...] = (),
    empty_text: str = "N/A",
) -> None:
    _heading(doc, title, level=level)
    if not rows:
        _paragraph(doc, empty_text, style="Body Text")
        return
    _table(doc, headers, rows, col_widths=widths)


def _render_document_cover(
    doc: DocClass,
    classification: DocumentationClassification,
    units: list[DocumentUnit],
) -> None:
    _centered_title(doc, "Functional Specification")
    scope_units = ", ".join(unit.unit_code for unit in units) or "All detected units"
    _paragraph(doc, f"Generated from source code for {scope_units}.")
    author_rows = [["Generated by SattLint", "", ""]]
    document_rows = [["", "Functional Specification", scope_units]]
    _table(doc, ["NNE Author", "NNE Author", "NNE Author"], author_rows, col_widths=(2, 2, 2))
    _table(doc, ["", "Document", "NN Doc. no."], document_rows, col_widths=(1, 3, 2))
    doc.add_page_break()


def _render_introduction(
    doc: DocClass,
    units: list[DocumentUnit],
) -> None:
    _heading(doc, "Introduction", level=1)

    _heading(doc, "Scope", level=2)
    _paragraph(
        doc,
        "This document is the generated Functional Specification (FS) for the detected process units and their configured physical and procedural elements.",
    )
    if units:
        _paragraph(doc, "The FS covers the following detected units:")
        for unit in units:
            _bullet(doc, f"{unit.unit_code} {unit.title}")

    _heading(doc, "S88 Model", level=2)
    _paragraph(doc, "The ANSI/ISA-88 (S88) model is used to structure the generated specification.")
    _underlined(doc, "S88 Physical Model")
    _paragraph(
        doc,
        "The physical model section groups units, equipment modules, device interfaces, and communication links detected from the source structure.",
    )
    _underlined(doc, "S88 Procedural Model")
    _paragraph(
        doc,
        "The procedural model section groups detected operations together with their parameters, messages, events, and available sequence logic.",
    )

    _heading(doc, "Definitions and abbreviations", level=2)
    _heading(doc, "Abbreviations", level=3)
    for abbreviation, meaning in (
        ("EM", "Equipment Module"),
        ("EP", "Engineering Parameter"),
        ("FS", "Functional Specification"),
        ("OP", "Operation"),
        ("RP", "Recipe Parameter"),
        ("UP", "User Parameter"),
    ):
        _paragraph(doc, f"{abbreviation} {meaning}")


def _render_references(
    doc: DocClass,
    classification: DocumentationClassification,
    units: list[DocumentUnit],
) -> None:
    _heading(doc, "References", level=2)
    _bullet(doc, f"Detected units: {len(units)}")
    _bullet(doc, f"Equipment modules: {sum(len(unit.equipment_modules) for unit in units)}")
    _bullet(doc, f"Operations: {sum(len(unit.operations) for unit in units)}")
    _bullet(doc, f"Recipe parameters: {sum(len(unit.recipe_parameters) for unit in units)}")
    _bullet(doc, f"Engineering parameters: {sum(len(unit.engineering_parameters) for unit in units)}")
    _bullet(doc, f"Uncategorized modules: {len(classification.uncategorized)}")


def _render_physical_model(
    doc: DocClass,
    units: list[DocumentUnit],
    classification: DocumentationClassification,
) -> None:
    _heading(doc, "S88 Physical model", level=1)
    process_cell_name = (
        ", ".join(sorted({unit.section_name for unit in units if unit.section_name})) or "Detected process cell"
    )
    _heading(doc, f"Process Cell - {process_cell_name}", level=2)
    _paragraph(doc, "The detected process cell comprises the units listed below.")

    unit_rows = [[unit.unit_code, unit.unit_class, unit.title, f"See section {unit.title}"] for unit in units]
    _table(doc, ["Unit", "Unit Class", "Danish Description", "Unit Definition"], unit_rows, col_widths=(1, 2, 2, 2))

    _render_named_table_section(
        doc,
        "Instance-configurable parameters",
        [row for unit in units for row in _configurable_parameter_rows(unit)],
        headers=["Unit Parameter", "Description", "Unit Instance", "Configured to"],
        level=3,
        widths=(2, 3, 1, 2),
    )

    for index, unit in enumerate(units, start=1):
        _render_unit_physical_section(doc, unit, classification, index=index)


def _render_unit_physical_section(
    doc: DocClass,
    unit: DocumentUnit,
    classification: DocumentationClassification,
    *,
    index: int,
) -> None:
    support_entries = _support_entries(unit, classification)

    _heading(doc, f"Unit Class definition: {unit.title}", level=2)
    _paragraph(
        doc, f"This section summarizes the detected physical structure for unit {unit.unit_code}.", style="Body Text"
    )
    _paragraph(doc, f"Overall detected unit class: {_prettify_name(unit.unit_class)}.", style="Caption")
    _paragraph(
        doc,
        "Note: The generated section is based on code structure, parameter mappings, and detected module groupings.",
    )

    _render_named_table_section(
        doc,
        "Equipment module instances",
        _module_list_rows(unit.equipment_modules),
        headers=["Equipment Module Name", "Description", "Definition"],
        level=3,
        widths=(2, 3, 2),
    )

    _render_named_table_section(
        doc,
        "Measurements and logging",
        _measurement_rows([entry for entry in support_entries if _is_measurement_entry(entry)]),
        headers=[
            "Measurement",
            "Tag",
            "Min",
            "Max",
            "Eng. Unit",
            "Log interval\n(Max)",
            "Dead-band\nrelative",
            "Log interval \n(Min)",
        ],
        level=3,
        widths=(2, 1, 1, 1, 1, 1, 1, 1),
    )
    _render_named_table_section(
        doc,
        "Special Logging",
        [
            [
                _module_description(entry),
                entry.name,
                _entry_variable_text(entry, "min", "lowlimit"),
                _entry_variable_text(entry, "max", "highlimit"),
                _first_non_empty(
                    _metadata_value(entry, "EngUnit", "Unit"), _entry_variable_text(entry, "engunit", "unit")
                ),
            ]
            for entry in support_entries
            if _is_special_logging_entry(entry)
        ],
        headers=["Log description", "Tag", "Min", "Max", "Eng. Unit"],
        level=3,
        widths=(3, 1, 1, 1, 1),
    )
    _render_named_table_section(
        doc,
        "Inlet Consumption Logging",
        [
            [
                _moduletype_summary(entry),
                _module_description(entry),
                entry.name,
                _metadata_value(entry, "Format"),
                _metadata_value(entry, "Measurement"),
                _first_non_empty(_metadata_value(entry, "Enabled"), "Yes"),
            ]
            for entry in support_entries
            if _is_inlet_consumption_entry(entry)
        ],
        headers=["Journal Type", "Description", "Log Parameter Name", "Format", "Measurement", "Enabled"],
        level=3,
        widths=(2, 3, 2, 1, 1, 1),
    )
    _render_named_table_section(
        doc,
        "Other Devices",
        _simple_name_tag_rows([entry for entry in support_entries if _is_other_device_entry(entry)], "description"),
        headers=["Device Description", "Tag"],
        level=3,
        widths=(4, 2),
    )
    _render_named_table_section(
        doc,
        "Timers",
        [
            [
                entry.name,
                _module_description(entry),
                _metadata_value(entry, "Operation"),
                _metadata_value(entry, "Operation2"),
            ]
            for entry in support_entries
            if _is_timer_entry(entry)
        ],
        headers=["Timer", "Description", "Operation", "Operation"],
        level=3,
        widths=(2, 3, 1, 1),
    )
    _render_named_table_section(
        doc,
        "Unit Events",
        _event_table_rows([entry for entry in support_entries if _is_event_entry(entry)]),
        headers=["Tag", "Event Text", "Sev.", "Activation"],
        level=3,
        widths=(2, 3, 1, 2),
    )
    _render_named_table_section(
        doc,
        "Interlocks",
        _interlock_rows([variable for variable in unit.root.localvariables if "interlock" in variable.name.casefold()]),
        headers=[
            "Tag",
            "Description",
            "Interlock Activation\n(Condition to cause interlock)",
            "Interlock Enable\n(Manipulation of equipment)",
        ],
        level=3,
        widths=(2, 2, 2, 2),
    )
    _render_named_table_section(
        doc,
        "Interventions",
        _generic_section_rows([entry for entry in support_entries if _is_intervention_entry(entry)]),
        headers=["Name", "Module Type", "Location", "Notes"],
        level=3,
        widths=(2, 2, 3, 2),
    )
    _render_named_table_section(
        doc,
        "Exceptions",
        _exception_rows(
            [variable for variable in unit.root.moduleparameters if variable.name.casefold().startswith("error_")]
        ),
        headers=["Condition", "Exceptions Description"],
        level=3,
        widths=(2, 4),
    )
    _render_named_table_section(
        doc,
        "Supervision",
        _generic_section_rows([entry for entry in support_entries if _is_supervision_entry(entry)]),
        headers=["Name", "Module Type", "Location", "Notes"],
        level=3,
        widths=(2, 2, 3, 2),
    )
    _render_named_table_section(
        doc,
        "Unit graphics",
        _generic_section_rows([entry for entry in support_entries if _is_graphics_entry(entry)]),
        headers=["Name", "Module Type", "Location", "Notes"],
        level=3,
        widths=(2, 2, 3, 2),
    )
    _render_named_table_section(
        doc,
        "CIP Flip of valves",
        _generic_section_rows([entry for entry in support_entries if _is_cip_valve_entry(entry)]),
        headers=["Name", "Module Type", "Location", "Notes"],
        level=3,
        widths=(2, 2, 3, 2),
    )
    _render_named_table_section(
        doc,
        "Calculations",
        _calculation_rows([entry for entry in support_entries if _is_calculation_entry(entry)]),
        headers=["Tag", "Description", "Calculation"],
        level=3,
        widths=(2, 3, 2),
    )
    _heading(doc, "Communication", level=3)
    from_rows = _communication_rows(unit, direction="from")
    to_rows = _communication_rows(unit, direction="to")
    if from_rows:
        _table(doc, ["From", "Comment"], from_rows, col_widths=(2, 4))
    else:
        _paragraph(doc, "N/A", style="Body Text")
    if to_rows:
        _table(doc, ["To", "Comment"], to_rows, col_widths=(2, 4))

    for equipment_module in unit.equipment_modules:
        _render_equipment_module_section(doc, equipment_module, classification)


def _render_equipment_module_section(
    doc: DocClass,
    entry: DocumentedModule,
    classification: DocumentationClassification,
) -> None:
    title = _module_title(entry)
    _heading(doc, f"Equipment Module {title}", level=3)
    _underlined(doc, "Description")
    _paragraph(doc, _module_description(entry))

    _render_named_table_section(
        doc,
        "States",
        _state_rows(entry, classification),
        headers=["State", "Description", "Logic"],
        level=4,
        widths=(2, 3, 2),
    )

    parameter_entries = classification.descendants(entry, category="ep")
    parameter_rows = _parameter_catalog_rows(parameter_entries)
    if not parameter_rows:
        parameter_rows = _variable_rows(
            [
                variable
                for variable in (*entry.moduleparameters, *entry.localvariables)
                if variable.name.casefold() not in {"name", "p", "colours"}
            ]
        )
        _render_named_table_section(
            doc,
            "Parameters",
            parameter_rows,
            headers=["Tag", "Datatype", "Init", "Description"],
            level=4,
            widths=(2, 2, 1, 3),
        )
    else:
        _render_named_table_section(
            doc,
            "Parameters",
            parameter_rows,
            headers=["Parameter", "Module Type", "Location", "Description"],
            level=4,
            widths=(2, 2, 3, 2),
        )

    _render_named_table_section(
        doc,
        "PID Controllers",
        _pid_controller_rows(entry, classification),
        headers=["Tag", "Module Type", "Location", "Description"],
        level=4,
        widths=(2, 2, 3, 2),
    )


def _render_procedural_model(
    doc: DocClass,
    units: list[DocumentUnit],
    classification: DocumentationClassification,
) -> None:
    _heading(doc, "S88 Procedural model", level=1)

    for unit in units:
        if not unit.operations and not unit.engineering_parameters:
            continue

        _heading(doc, f"Operations Unit Class - {unit.title}", level=2)
        _render_named_table_section(
            doc,
            "Unit engineering parameters",
            _parameter_catalog_rows(unit.engineering_parameters),
            headers=["Parameter", "Module Type", "Location", "Description"],
            level=3,
            widths=(2, 2, 3, 2),
        )

        for operation in unit.operations:
            _render_operation_section(doc, operation, classification)


def _render_operation_section(
    doc: DocClass,
    entry: DocumentedModule,
    classification: DocumentationClassification,
) -> None:
    title = _module_title(entry)
    _heading(doc, f"Operation {title}", level=3)
    _underlined(doc, "Description")
    _paragraph(doc, _module_description(entry))

    recipe_parameters = classification.descendants(entry, category="rp")
    engineering_parameters = classification.descendants(entry, category="ep")
    user_parameters = classification.descendants(entry, category="up")

    _render_named_table_section(
        doc,
        "Recipe parameters",
        _parameter_catalog_rows(recipe_parameters),
        headers=["Parameter", "Module Type", "Location", "Description"],
        level=4,
        widths=(2, 2, 3, 2),
    )
    _render_named_table_section(
        doc,
        "Engineering parameters",
        _parameter_catalog_rows(engineering_parameters),
        headers=["Parameter", "Module Type", "Location", "Description"],
        level=4,
        widths=(2, 2, 3, 2),
    )
    if user_parameters:
        _render_named_table_section(
            doc,
            "User parameters",
            _parameter_catalog_rows(user_parameters),
            headers=["Parameter", "Module Type", "Location", "Description"],
            level=4,
            widths=(2, 2, 3, 2),
        )

    _render_named_table_section(
        doc,
        "Messages",
        _message_rows(entry, classification),
        headers=["Name", "Module Type", "Location", "Description"],
        level=4,
        widths=(2, 2, 3, 2),
    )
    _render_named_table_section(
        doc,
        "Operation events",
        _event_rows(entry, classification),
        headers=["Name", "Module Type", "Location", "Description"],
        level=4,
        widths=(2, 2, 3, 2),
    )
    _render_named_table_section(
        doc,
        "Special log parameters",
        _special_logging_rows(entry, classification),
        headers=["Name", "Module Type", "Location", "Description"],
        level=4,
        widths=(2, 2, 3, 2),
    )

    for sequence_name, sequence_rows in _sequence_rows(entry, classification):
        _heading(doc, f"Sub sequence - {sequence_name}", level=4)
        if sequence_rows:
            _table(
                doc,
                ["Type", "Name", "Condition / Detail", "Enter", "Active", "Exit"],
                _sequence_table_rows(sequence_rows),
                col_widths=(1, 1, 3, 2, 2, 2),
            )
        else:
            _paragraph(doc, "No explicit sequence statements detected.")


def _render_uncategorized_appendix(
    doc: DocClass,
    classification: DocumentationClassification,
    units: list[DocumentUnit],
) -> None:
    unit_roots = [unit.root for unit in units]
    appendix_entries = [entry for entry in classification.uncategorized if not _is_within_any(entry, unit_roots)]
    if not appendix_entries:
        return

    _heading(doc, "Appendix: Supporting modules", level=1)
    _render_named_table_section(
        doc,
        "Other module instances",
        _generic_section_rows(appendix_entries),
        headers=["Name", "Module Type", "Location", "Description"],
        level=2,
        widths=(2, 2, 3, 2),
    )


def _render_upgrade_insights(
    doc: DocClass,
    upgrade_issues: t.Sequence[Issue],
) -> None:
    if not upgrade_issues:
        return

    _heading(doc, "Upgrade insights", level=1)
    _paragraph(
        doc,
        "Repeated module names with structural drift are summarized below to support upgrade planning and regression review.",
    )

    for issue in upgrade_issues:
        issue_data = issue.data or {}
        module_name = str(issue_data.get("module_name", issue.message))
        _heading(doc, f"Module {module_name}", level=2)
        _paragraph(doc, issue.message)
        for note in issue_data.get("upgrade_notes", []) or []:
            _bullet(doc, str(note))


def _render_change_log(doc: DocClass) -> None:
    _heading(doc, "Change Log", level=1)
    _table(
        doc,
        ["Revision", "Description"],
        [["Generated", "Document generated from SattLine source structure by SattLint."]],
        col_widths=(1, 5),
    )


def _variable_rows(variables: list[Variable]) -> list[list[str]]:
    return [
        [
            variable.name,
            variable.datatype_text,
            _value_text(variable.init_value),
            variable.description or "",
        ]
        for variable in variables
    ]


def _is_within_any(entry: DocumentedModule, ancestors: list[DocumentedModule]) -> bool:
    return any(
        entry.path != ancestor.path and entry.path[: len(ancestor.path)] == ancestor.path for ancestor in ancestors
    )


def _is_measurement_entry(entry: DocumentedModule) -> bool:
    label_cf = _moduletype_summary(entry).casefold()
    return bool(
        _TAG_NAME_RE.match(entry.name)
        or _DEVICE_PREFIX_RE.match(entry.name)
        or any(token in label_cf for token in ("analog", "switch", "flow", "scale", "pressure"))
    )


def _is_special_logging_entry(entry: DocumentedModule) -> bool:
    label_cf = _moduletype_summary(entry).casefold()
    name_cf = entry.name.casefold()
    return "journal" in label_cf or ("log" in name_cf and "inletcons" not in name_cf)


def _is_inlet_consumption_entry(entry: DocumentedModule) -> bool:
    label_cf = _moduletype_summary(entry).casefold()
    name_cf = entry.name.casefold()
    return "inletcons" in name_cf or "incons" in label_cf


def _is_timer_entry(entry: DocumentedModule) -> bool:
    label_cf = _moduletype_summary(entry).casefold()
    return "timer" in entry.name.casefold() or "timer" in label_cf


def _is_event_entry(entry: DocumentedModule) -> bool:
    label_cf = _moduletype_summary(entry).casefold()
    name_cf = entry.name.casefold()
    return name_cf.startswith("event") or "event" in label_cf or "journal" in label_cf


def _is_graphics_entry(entry: DocumentedModule) -> bool:
    label_cf = _moduletype_summary(entry).casefold()
    return any(token in label_cf for token in ("icon", "view", "toggle", "header", "infotext", "transferdisplay"))


def _is_calculation_entry(entry: DocumentedModule) -> bool:
    label_cf = _moduletype_summary(entry).casefold()
    return any(token in label_cf for token in ("minmax", "pid", "realtoreal", "calc", "ctrl"))


def _is_cip_valve_entry(entry: DocumentedModule) -> bool:
    label_cf = _moduletype_summary(entry).casefold()
    name_cf = entry.name.casefold()
    return name_cf.startswith("v") and "cip" in label_cf


def _is_other_device_entry(entry: DocumentedModule) -> bool:
    if _is_measurement_entry(entry) or _is_graphics_entry(entry) or _is_special_logging_entry(entry):
        return False
    return _DEVICE_PREFIX_RE.match(entry.name) is not None


def _is_intervention_entry(entry: DocumentedModule) -> bool:
    label_cf = _moduletype_summary(entry).casefold()
    return "button" in label_cf or "manual" in label_cf


def _is_supervision_entry(entry: DocumentedModule) -> bool:
    label_cf = _moduletype_summary(entry).casefold()
    name_cf = entry.name.casefold()
    return any(token in label_cf or token in name_cf for token in ("warning", "alarm", "supervision"))


def generate_docx(
    root,
    out_path: str | pathlib.Path,
    *,
    documentation_config: dict | None = None,
    unavailable_libraries: set[str] | None = None,
    upgrade_issues: t.Sequence[Issue] | None = None,
) -> None:
    doc: DocClass = DocumentFactory()
    _ensure_styles(doc)
    classification = classify_documentation_structure(
        root,
        documentation_config or config_module.get_documentation_config(),
        unavailable_libraries=unavailable_libraries,
    )
    units = _build_units(classification)

    _render_document_cover(doc, classification, units)
    _render_introduction(doc, units)
    _render_references(doc, classification, units)
    _render_physical_model(doc, units, classification)
    doc.add_page_break()
    _render_procedural_model(doc, units, classification)
    _render_uncategorized_appendix(doc, classification, units)
    if upgrade_issues is None:
        from ..analyzers.modules import analyze_version_drift

        upgrade_issues = analyze_version_drift(root).issues
    resolved_upgrade_issues: t.Sequence[Issue] = upgrade_issues or ()
    _render_upgrade_insights(doc, resolved_upgrade_issues)
    _render_change_log(doc)

    out_file = pathlib.Path(out_path)
    doc.save(str(out_file))
    log.info("Documentation written to %s", out_file.resolve())

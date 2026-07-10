from __future__ import annotations

import typing as t
from dataclasses import dataclass

from docx.document import Document as DocClass
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from sattline_parser.models.ast_model import (
    Sequence,
    SFCAlternative,
    SFCBreak,
    SFCFork,
    SFCParallel,
    SFCStep,
    SFCSubsequence,
    SFCTransition,
    SFCTransitionSub,
)
from sattline_parser.utils.formatter import format_expr

from ..classification import DocumentedModule


@dataclass(frozen=True)
class DocumentUnit:
    root: DocumentedModule
    unit_code: str
    title: str
    unit_class: str
    section_name: str
    equipment_modules: list[DocumentedModule]
    operations: list[DocumentedModule]
    recipe_parameters: list[DocumentedModule]
    engineering_parameters: list[DocumentedModule]
    user_parameters: list[DocumentedModule]


@dataclass(frozen=True)
class SequenceRenderRow:
    node_type: str
    name: str
    detail: str
    enter: str = ""
    active: str = ""
    exit: str = ""


def _ensure_styles(doc: DocClass) -> None:
    styles = t.cast(t.Any, doc.styles)
    if "Underlined" not in styles:
        style = styles.add_style("Underlined", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.underline = True
        style.font.size = Pt(10)


def _heading(doc: DocClass, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _paragraph(
    doc: DocClass,
    text: str,
    *,
    bold: bool = False,
    style: str | None = None,
) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.bold = bold


def _centered_title(doc: DocClass, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def _bullet(doc: DocClass, text: str) -> None:
    doc.add_paragraph(text, style="List Paragraph")


def _underlined(doc: DocClass, text: str) -> None:
    _paragraph(doc, text, style="Underlined")


def _table(
    doc: DocClass,
    headers: list[str],
    rows: t.Sequence[t.Sequence[object]],
    col_widths: tuple[int, ...] = (),
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)

    if col_widths:
        for index, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[index].width = Inches(width)


def _value_text(value: object | None) -> str:
    if value is None:
        return ""
    return str(value)


def _format_coord(value: tuple[float, ...] | None) -> str:
    if value is None:
        return "<none>"
    return f"({', '.join(f'{part:g}' for part in value)})"


def _prettify_name(text: str) -> str:
    if not text:
        return ""
    normalized = text.replace("_", " ").strip()
    if normalized.isupper():
        return normalized
    return normalized


def _sequence_code_text(statements: list[object]) -> str:
    return "\n".join(format_expr(statement) for statement in statements)


def _append_sequence_rows(
    rows: list[SequenceRenderRow],
    nodes: list[object],
    *,
    context: str | None = None,
) -> None:
    for node in nodes:
        detail_prefix = f"{context}: " if context else ""
        if isinstance(node, SFCStep):
            rows.append(
                SequenceRenderRow(
                    node_type="Init step" if node.kind == "init" else "Step",
                    name=node.name,
                    detail=detail_prefix + ("Initial step" if node.kind == "init" else "Execution step"),
                    enter=_sequence_code_text(node.code.enter),
                    active=_sequence_code_text(node.code.active),
                    exit=_sequence_code_text(node.code.exit),
                )
            )
            continue

        if isinstance(node, SFCTransition):
            rows.append(
                SequenceRenderRow(
                    node_type="Transition",
                    name=node.name or "<unnamed>",
                    detail=detail_prefix + format_expr(node.condition),
                )
            )
            continue

        if isinstance(node, SFCAlternative):
            rows.append(
                SequenceRenderRow(
                    node_type="Alternative",
                    name="",
                    detail=detail_prefix + f"{len(node.branches)} branches",
                )
            )
            for index, branch in enumerate(node.branches, start=1):
                branch_context = f"Alternative branch {index}"
                rows.append(
                    SequenceRenderRow(
                        node_type="Branch",
                        name=f"Branch {index}",
                        detail=detail_prefix + "Alternative path",
                    )
                )
                _append_sequence_rows(rows, branch, context=branch_context)
            continue

        if isinstance(node, SFCParallel):
            rows.append(
                SequenceRenderRow(
                    node_type="Parallel",
                    name="",
                    detail=detail_prefix + f"{len(node.branches)} branches",
                )
            )
            for index, branch in enumerate(node.branches, start=1):
                branch_context = f"Parallel branch {index}"
                rows.append(
                    SequenceRenderRow(
                        node_type="Branch",
                        name=f"Branch {index}",
                        detail=detail_prefix + "Parallel path",
                    )
                )
                _append_sequence_rows(rows, branch, context=branch_context)
            continue

        if isinstance(node, SFCSubsequence):
            rows.append(
                SequenceRenderRow(
                    node_type="Subsequence",
                    name=node.name,
                    detail=detail_prefix + "Nested sequence",
                )
            )
            _append_sequence_rows(rows, node.body, context=f"Subsequence {node.name}")
            continue

        if isinstance(node, SFCTransitionSub):
            rows.append(
                SequenceRenderRow(
                    node_type="Transition section",
                    name=node.name,
                    detail=detail_prefix + "Nested transition sequence",
                )
            )
            _append_sequence_rows(rows, node.body, context=f"Transition section {node.name}")
            continue

        if isinstance(node, SFCFork):
            rows.append(
                SequenceRenderRow(
                    node_type="Fork",
                    name=", ".join(node.targets),
                    detail=detail_prefix + ("Fork targets" if len(node.targets) > 1 else "Fork target"),
                )
            )
            continue

        if isinstance(node, SFCBreak):
            rows.append(
                SequenceRenderRow(
                    node_type="Break",
                    name="",
                    detail=detail_prefix + "Break sequence flow",
                )
            )
            continue

        rows.append(
            SequenceRenderRow(
                node_type="Statement",
                name="",
                detail=detail_prefix + format_expr(node),
            )
        )


def _sequence_render_rows(sequence: Sequence) -> list[SequenceRenderRow]:
    rows: list[SequenceRenderRow] = []
    _append_sequence_rows(rows, list(sequence.code or []))
    return rows


def _sequence_table_rows(rows: list[SequenceRenderRow]) -> list[list[str]]:
    return [[row.node_type, row.name, row.detail, row.enter, row.active, row.exit] for row in rows]
